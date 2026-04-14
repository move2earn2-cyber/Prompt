#!/usr/bin/env python3
"""
falcon_perf_assessment.py  (v2)
Performance Impact Assessment for CrowdStrike Falcon Operator
  - FalconNodeSensor (DaemonSet)
  - FalconAdmission / Falcon KAC (Admission Controller)

Usage:
  # Step 1 – baseline BEFORE deploying Falcon
  python3 falcon_perf_assessment.py --mode baseline --output /tmp/falcon_baseline.json

  # Step 2 – post-deploy report (HTML + DOCX)
  python3 falcon_perf_assessment.py --mode post \
      --baseline /tmp/falcon_baseline.json \
      --output /tmp/falcon_report

  # Continuous sampling
  python3 falcon_perf_assessment.py --mode sample \
      --samples 10 --interval 30 --output /tmp/falcon_samples.json

Dependencies:
  pip install kubernetes jinja2 plotly kaleido python-docx
"""

import argparse
import json
import subprocess
import sys
import time
import datetime
import os
import io
from pathlib import Path

try:
    from kubernetes import client, config
    HAS_K8S_CLIENT = True
except ImportError:
    HAS_K8S_CLIENT = False

try:
    import plotly.graph_objects as go
    from plotly.subplots import make_subplots
    HAS_PLOTLY = True
except ImportError:
    HAS_PLOTLY = False

try:
    from jinja2 import Template
    HAS_JINJA = True
except ImportError:
    HAS_JINJA = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ─────────────────────────────────────────────────────────────────────────────
# CONSTANTS
# ─────────────────────────────────────────────────────────────────────────────
FALCON_NAMESPACES       = ["falcon-system", "falcon-operator", "falcon-kac", "crowdstrike"]
FALCON_NODE_SENSOR_DS   = "falcon-node-sensor"
FALCON_ADMISSION_DEPLOY = "falcon-admission"
WEBHOOK_NAME_PATTERN    = "falcon"


# ─────────────────────────────────────────────────────────────────────────────
# KUBECTL HELPERS
# ─────────────────────────────────────────────────────────────────────────────
def run_kubectl(args: list, timeout: int = 30):
    cmd = ["kubectl"] + args
    try:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=timeout)
        return result.stdout, result.stderr, result.returncode
    except subprocess.TimeoutExpired:
        return "", f"kubectl timeout after {timeout}s", 1
    except FileNotFoundError:
        return "", "kubectl not found in PATH", 1


def kubectl_json(args: list):
    stdout, stderr, rc = run_kubectl(args + ["-o", "json"])
    if rc != 0:
        print(f"  [WARN] kubectl error: {stderr.strip()}", file=sys.stderr)
        return None
    try:
        return json.loads(stdout)
    except json.JSONDecodeError:
        return None


def kubectl_raw_metrics(path: str) -> str:
    stdout, _, _ = run_kubectl(["get", "--raw", path])
    return stdout


# ─────────────────────────────────────────────────────────────────────────────
# CLUSTER IDENTITY
# ─────────────────────────────────────────────────────────────────────────────
def collect_cluster_info() -> dict:
    """
    Return cluster name, server URL, context, and Kubernetes version.
    Sources: kubectl config current-context, kubectl version, kubectl cluster-info.
    """
    info = {
        "cluster_name":    "unknown",
        "context":         "unknown",
        "server_url":      "unknown",
        "k8s_version":     "unknown",
        "k8s_platform":    "unknown",
    }

    # Current context name
    stdout, _, rc = run_kubectl(["config", "current-context"])
    if rc == 0:
        info["context"] = stdout.strip()

    # Cluster name from kubeconfig (may differ from context)
    cfg = kubectl_json(["config", "view"])
    if cfg:
        current_ctx = info["context"]
        # Find matching context entry
        for ctx_entry in cfg.get("contexts", []):
            if ctx_entry.get("name") == current_ctx:
                cluster_ref = ctx_entry.get("context", {}).get("cluster", "")
                info["cluster_name"] = cluster_ref or current_ctx
                break
        # Find matching cluster server URL
        for cl_entry in cfg.get("clusters", []):
            if cl_entry.get("name") == info["cluster_name"]:
                info["server_url"] = cl_entry.get("cluster", {}).get("server", "unknown")
                break

    # Kubernetes version
    ver = kubectl_json(["version"])
    if ver:
        sv = ver.get("serverVersion", {})
        info["k8s_version"] = sv.get("gitVersion", "unknown")
        # Detect managed platform from version string or git version
        gv = sv.get("gitVersion", "").lower()
        if "eks" in gv:
            info["k8s_platform"] = "Amazon EKS"
        elif "gke" in gv:
            info["k8s_platform"] = "Google GKE"
        elif "aks" in gv:
            info["k8s_platform"] = "Azure AKS"
        elif "openshift" in gv or "okd" in gv:
            info["k8s_platform"] = "OpenShift"
        else:
            # Try node labels for platform hints
            nodes = kubectl_json(["get", "nodes"])
            if nodes:
                for n in nodes.get("items", []):
                    labels = n.get("metadata", {}).get("labels", {})
                    pv_id = n.get("spec", {}).get("providerID", "")
                    if "eks" in pv_id:
                        info["k8s_platform"] = "Amazon EKS"; break
                    elif "gce" in pv_id:
                        info["k8s_platform"] = "Google GKE"; break
                    elif "azure" in pv_id:
                        info["k8s_platform"] = "Azure AKS"; break
                    elif labels.get("node.openshift.io/os_id"):
                        info["k8s_platform"] = "OpenShift"; break
                else:
                    info["k8s_platform"] = "Vanilla / On-Prem"

    # Node count
    nodes = kubectl_json(["get", "nodes"])
    if nodes:
        items = nodes.get("items", [])
        info["node_count"] = len(items)
        info["worker_count"] = sum(
            1 for n in items
            if not n.get("metadata", {}).get("labels", {}).get("node-role.kubernetes.io/control-plane")
            and not n.get("metadata", {}).get("labels", {}).get("node-role.kubernetes.io/master")
        )

    return info


# ─────────────────────────────────────────────────────────────────────────────
# METRIC COLLECTION  (unchanged from v1)
# ─────────────────────────────────────────────────────────────────────────────
def collect_node_resources() -> list:
    stdout, stderr, rc = run_kubectl(["top", "nodes", "--no-headers"])
    nodes = []
    if rc != 0:
        print(f"  [WARN] top nodes failed: {stderr.strip()}", file=sys.stderr)
        return nodes
    for line in stdout.strip().splitlines():
        parts = line.split()
        if len(parts) < 5:
            continue
        nodes.append({
            "name":      parts[0],
            "cpu_cores": _parse_cpu(parts[1]),
            "mem_mib":   _parse_mem(parts[3]),
        })
    return nodes


def collect_pod_resources(namespace: str = "") -> list:
    args = ["top", "pods", "--no-headers"]
    if namespace:
        args += ["-n", namespace]
    else:
        args += ["-A"]
    stdout, stderr, rc = run_kubectl(args)
    pods = []
    if rc != 0:
        print(f"  [WARN] top pods failed: {stderr.strip()}", file=sys.stderr)
        return pods
    for line in stdout.strip().splitlines():
        parts = line.split()
        if len(parts) < 4:
            continue
        if namespace:
            ns, pod, cpu_raw, mem_raw = namespace, parts[0], parts[1], parts[2]
        else:
            ns, pod, cpu_raw, mem_raw = parts[0], parts[1], parts[2], parts[3]
        pods.append({
            "namespace": ns,
            "pod":       pod,
            "cpu_cores": _parse_cpu(cpu_raw),
            "mem_mib":   _parse_mem(mem_raw),
        })
    return pods


def collect_falcon_pods() -> list:
    all_pods = collect_pod_resources()
    return [
        p for p in all_pods
        if any(ns in p["namespace"] for ns in FALCON_NAMESPACES)
        or "falcon" in p["pod"].lower()
        or "crowdstrike" in p["pod"].lower()
    ]


def collect_daemonset_status() -> dict:
    for ns in FALCON_NAMESPACES:
        ds = kubectl_json(["get", "daemonset", FALCON_NODE_SENSOR_DS, "-n", ns])
        if ds:
            status = ds.get("status", {})
            return {
                "namespace":     ns,
                "desired":       status.get("desiredNumberScheduled", 0),
                "ready":         status.get("numberReady", 0),
                "available":     status.get("numberAvailable", 0),
                "misscheduled":  status.get("numberMisscheduled", 0),
                "update_strategy": ds.get("spec", {}).get("updateStrategy", {}).get("type", "unknown"),
            }
    return {}


def collect_admission_status() -> dict:
    for ns in FALCON_NAMESPACES:
        for deploy_name in [FALCON_ADMISSION_DEPLOY, "falcon-kac", "falcon-kac-controller"]:
            dep = kubectl_json(["get", "deployment", deploy_name, "-n", ns])
            if dep:
                status = dep.get("status", {})
                spec   = dep.get("spec", {})
                return {
                    "namespace":  ns,
                    "name":       deploy_name,
                    "desired":    spec.get("replicas", 0),
                    "ready":      status.get("readyReplicas", 0),
                    "available":  status.get("availableReplicas", 0),
                    "conditions": [
                        {"type": c["type"], "status": c["status"]}
                        for c in status.get("conditions", [])
                    ],
                }
    return {}


def collect_webhook_config() -> dict:
    vwcs = kubectl_json(["get", "validatingwebhookconfigurations"])
    if not vwcs:
        return {}
    for item in vwcs.get("items", []):
        name = item.get("metadata", {}).get("name", "")
        if WEBHOOK_NAME_PATTERN in name.lower():
            webhooks = item.get("webhooks", [])
            wh_info = []
            for wh in webhooks:
                wh_info.append({
                    "name":           wh.get("name", ""),
                    "failurePolicy":  wh.get("failurePolicy", "unknown"),
                    "timeoutSeconds": wh.get("timeoutSeconds", "default"),
                    "matchPolicy":    wh.get("matchPolicy", "Equivalent"),
                    "namespaceSelector": bool(wh.get("namespaceSelector")),
                    "rules":          len(wh.get("rules", [])),
                })
            return {"name": name, "webhooks": wh_info}
    return {}


def collect_admission_latency_from_metrics() -> dict:
    raw = kubectl_raw_metrics("/metrics")
    latency = {"p50": None, "p95": None, "p99": None, "sample_count": 0}
    if not raw:
        return latency
    buckets = {}
    total_count, total_sum = 0, 0.0
    for line in raw.splitlines():
        if "apiserver_admission_webhook_admission_duration_seconds" not in line:
            continue
        if WEBHOOK_NAME_PATTERN not in line.lower():
            continue
        if line.startswith("#"):
            continue
        if "_bucket{" in line:
            try:
                le_start = line.index('le="') + 4
                le_end   = line.index('"', le_start)
                le_val   = line[le_start:le_end]
                val      = float(line.split()[-1])
                if le_val != "+Inf":
                    buckets[float(le_val)] = val
            except (ValueError, IndexError):
                pass
        elif "_count{" in line:
            try:
                total_count = int(float(line.split()[-1]))
            except (ValueError, IndexError):
                pass
        elif "_sum{" in line:
            try:
                total_sum = float(line.split()[-1])
            except (ValueError, IndexError):
                pass
    if total_count > 0:
        latency["sample_count"] = total_count
        latency["mean_ms"] = round((total_sum / total_count) * 1000, 2)
    if buckets and total_count > 0:
        sorted_bounds = sorted(buckets.keys())
        for pct, label in [(0.5, "p50"), (0.95, "p95"), (0.99, "p99")]:
            target = pct * total_count
            for bound in sorted_bounds:
                if buckets[bound] >= target:
                    latency[label] = round(bound * 1000, 2)
                    break
    return latency


def collect_node_capacity() -> list:
    nodes_obj = kubectl_json(["get", "nodes"])
    if not nodes_obj:
        return []
    result = []
    for item in nodes_obj.get("items", []):
        name        = item["metadata"]["name"]
        allocatable = item.get("status", {}).get("allocatable", {})
        capacity    = item.get("status", {}).get("capacity", {})
        result.append({
            "name":               name,
            "allocatable_cpu":    _parse_cpu(allocatable.get("cpu", "0")),
            "allocatable_mem_mib": _parse_mem(allocatable.get("memory", "0Ki")),
            "capacity_cpu":       _parse_cpu(capacity.get("cpu", "0")),
            "capacity_mem_mib":   _parse_mem(capacity.get("memory", "0Ki")),
            "roles":              _get_node_roles(item),
        })
    return result


def collect_events() -> list:
    events = []
    for ns in FALCON_NAMESPACES:
        evts = kubectl_json(["get", "events", "-n", ns,
                             "--field-selector=type=Warning",
                             "--sort-by=.lastTimestamp"])
        if evts:
            for e in evts.get("items", [])[-10:]:
                events.append({
                    "namespace": ns,
                    "reason":   e.get("reason", ""),
                    "message":  e.get("message", "")[:120],
                    "count":    e.get("count", 1),
                    "last_ts":  e.get("lastTimestamp", ""),
                })
    return events


# ─────────────────────────────────────────────────────────────────────────────
# PARSING HELPERS
# ─────────────────────────────────────────────────────────────────────────────
def _parse_cpu(raw: str) -> float:
    raw = raw.strip()
    if raw.endswith("m"):
        return float(raw[:-1]) / 1000
    try:
        return float(raw)
    except ValueError:
        return 0.0


def _parse_mem(raw: str) -> float:
    raw = raw.strip()
    if raw.endswith("Ki"):
        return float(raw[:-2]) / 1024
    elif raw.endswith("Mi"):
        return float(raw[:-2])
    elif raw.endswith("Gi"):
        return float(raw[:-2]) * 1024
    elif raw.endswith("Ti"):
        return float(raw[:-2]) * 1024 * 1024
    elif raw.endswith("k") or raw.endswith("K"):
        return float(raw[:-1]) / 1024
    try:
        return float(raw) / (1024 * 1024)
    except ValueError:
        return 0.0


def _get_node_roles(node_item: dict) -> str:
    labels = node_item.get("metadata", {}).get("labels", {})
    roles  = [k.split("/")[-1] for k in labels if "node-role.kubernetes.io" in k]
    return ",".join(roles) if roles else "worker"


# ─────────────────────────────────────────────────────────────────────────────
# SNAPSHOT
# ─────────────────────────────────────────────────────────────────────────────
def collect_snapshot(label: str = "snapshot") -> dict:
    print(f"\n[{label.upper()}] Collecting metrics at {datetime.datetime.now().isoformat()}")
    snap = {
        "label":             label,
        "timestamp":         datetime.datetime.now().isoformat(),
        "cluster_info":      {},
        "node_resources":    [],
        "node_capacity":     [],
        "falcon_pods":       [],
        "all_pods":          [],
        "daemonset_status":  {},
        "admission_status":  {},
        "webhook_config":    {},
        "admission_latency": {},
        "apiserver_latency": {},
        "events":            [],
    }

    print("  → Cluster identity ...")
    snap["cluster_info"] = collect_cluster_info()
    print(f"     Cluster : {snap['cluster_info'].get('cluster_name')}")
    print(f"     Context : {snap['cluster_info'].get('context')}")
    print(f"     Platform: {snap['cluster_info'].get('k8s_platform')}")
    print(f"     Version : {snap['cluster_info'].get('k8s_version')}")

    print("  → Node resource usage ...")
    snap["node_resources"] = collect_node_resources()

    print("  → Node capacity/allocatable ...")
    snap["node_capacity"] = collect_node_capacity()

    print("  → All pod resource usage ...")
    snap["all_pods"] = collect_pod_resources()

    print("  → Falcon pod resources ...")
    snap["falcon_pods"] = collect_falcon_pods()

    print("  → FalconNodeSensor DaemonSet status ...")
    snap["daemonset_status"] = collect_daemonset_status()

    print("  → FalconAdmission deployment status ...")
    snap["admission_status"] = collect_admission_status()

    print("  → Webhook configuration ...")
    snap["webhook_config"] = collect_webhook_config()

    print("  → Admission webhook latency from /metrics ...")
    snap["admission_latency"] = collect_admission_latency_from_metrics()

    print("  → Warning events ...")
    snap["events"] = collect_events()

    return snap


def collect_samples(n: int, interval_s: int) -> list:
    samples = []
    for i in range(n):
        print(f"\n── Sample {i+1}/{n} ──────────────────────────────────────")
        samples.append(collect_snapshot(label=f"sample_{i+1}"))
        if i < n - 1:
            print(f"  Sleeping {interval_s}s ...")
            time.sleep(interval_s)
    return samples


# ─────────────────────────────────────────────────────────────────────────────
# DELTA / ANALYSIS
# ─────────────────────────────────────────────────────────────────────────────
def compute_delta(baseline: dict, post: dict) -> dict:
    delta = {"nodes": [], "summary": {}}
    base_nodes = {n["name"]: n for n in baseline.get("node_resources", [])}
    post_nodes = {n["name"]: n for n in post.get("node_resources", [])}
    for name, post_n in post_nodes.items():
        base_n = base_nodes.get(name, {"cpu_cores": 0, "mem_mib": 0})
        delta["nodes"].append({
            "name":          name,
            "cpu_delta":     round(post_n["cpu_cores"] - base_n["cpu_cores"], 4),
            "mem_delta_mib": round(post_n["mem_mib"]   - base_n["mem_mib"],   2),
            "cpu_base":      base_n["cpu_cores"],
            "cpu_post":      post_n["cpu_cores"],
            "mem_base":      base_n["mem_mib"],
            "mem_post":      post_n["mem_mib"],
        })
    total_base_cpu = sum(n["cpu_cores"] for n in baseline.get("node_resources", []))
    total_post_cpu = sum(n["cpu_cores"] for n in post.get("node_resources", []))
    total_base_mem = sum(n["mem_mib"]   for n in baseline.get("node_resources", []))
    total_post_mem = sum(n["mem_mib"]   for n in post.get("node_resources", []))
    falcon_cpu     = sum(p["cpu_cores"] for p in post.get("falcon_pods", []))
    falcon_mem     = sum(p["mem_mib"]   for p in post.get("falcon_pods", []))
    delta["summary"] = {
        "cluster_cpu_delta":     round(total_post_cpu - total_base_cpu, 4),
        "cluster_mem_delta_mib": round(total_post_mem - total_base_mem, 2),
        "falcon_cpu_cores":      round(falcon_cpu, 4),
        "falcon_mem_mib":        round(falcon_mem, 2),
        "cpu_pct_increase":      round((total_post_cpu - total_base_cpu) / max(total_base_cpu, 0.001) * 100, 2),
        "mem_pct_increase":      round((total_post_mem - total_base_mem) / max(total_base_mem, 0.001) * 100, 2),
        "cluster_cpu_base":      round(total_base_cpu, 4),
        "cluster_cpu_post":      round(total_post_cpu, 4),
        "cluster_mem_base":      round(total_base_mem, 2),
        "cluster_mem_post":      round(total_post_mem, 2),
    }
    base_lat = baseline.get("admission_latency", {})
    post_lat  = post.get("admission_latency", {})
    delta["admission_latency_delta"] = {
        "p50_delta_ms": _safe_diff(post_lat.get("p50"), base_lat.get("p50")),
        "p95_delta_ms": _safe_diff(post_lat.get("p95"), base_lat.get("p95")),
        "p99_delta_ms": _safe_diff(post_lat.get("p99"), base_lat.get("p99")),
    }
    return delta


def _safe_diff(a, b):
    if a is None or b is None:
        return None
    return round(a - b, 2)


def assess_risk(delta: dict, post: dict) -> dict:
    findings  = []
    risk_level = "LOW"
    summary   = delta.get("summary", {})
    cpu_pct   = summary.get("cpu_pct_increase", 0)
    mem_pct   = summary.get("mem_pct_increase", 0)

    if cpu_pct > 20:
        risk_level = "HIGH"
        findings.append(f"WARN  Cluster CPU increased by {cpu_pct}% — consider tuning sensor CPU limits.")
    elif cpu_pct > 10:
        risk_level = "MEDIUM"
        findings.append(f"INFO  Cluster CPU increased by {cpu_pct}% — monitor during peak load.")
    else:
        findings.append(f"OK    Cluster CPU overhead acceptable ({cpu_pct}%).")

    if mem_pct > 15:
        risk_level = "HIGH"
        findings.append(f"WARN  Cluster memory increased by {mem_pct}% — validate node memory headroom.")
    else:
        findings.append(f"OK    Cluster memory overhead acceptable ({mem_pct}%).")

    wh = post.get("webhook_config", {})
    for hook in wh.get("webhooks", []):
        timeout = hook.get("timeoutSeconds", 10)
        failure = hook.get("failurePolicy", "Fail")
        if failure == "Fail":
            findings.append(
                f"WARN  Webhook '{hook['name']}' has failurePolicy=Fail / timeout={timeout}s — "
                "KAC outage will block pod scheduling."
            )
            risk_level = "MEDIUM" if risk_level == "LOW" else risk_level
        else:
            findings.append(
                f"OK    Webhook '{hook['name']}' failurePolicy=Ignore — KAC failures won't block scheduling."
            )

    lat = post.get("admission_latency", {})
    p99 = lat.get("p99")
    if p99 and p99 > 1000:
        risk_level = "HIGH"
        findings.append(f"WARN  KAC webhook p99 latency {p99}ms — exceeds 1000ms threshold.")
    elif p99 and p99 > 500:
        findings.append(f"INFO  KAC webhook p99 latency {p99}ms — approaching 1000ms threshold.")
    elif p99:
        findings.append(f"OK    KAC webhook p99 latency {p99}ms — within safe range.")

    ds = post.get("daemonset_status", {})
    if ds:
        desired, ready = ds.get("desired", 0), ds.get("ready", 0)
        if desired > 0 and ready < desired:
            risk_level = "HIGH"
            findings.append(f"WARN  FalconNodeSensor: only {ready}/{desired} pods Ready.")
        elif desired > 0:
            findings.append(f"OK    FalconNodeSensor: {ready}/{desired} pods Ready.")

    return {"risk_level": risk_level, "findings": findings}


# ─────────────────────────────────────────────────────────────────────────────
# CHART GENERATION
# ─────────────────────────────────────────────────────────────────────────────
def generate_charts(baseline: dict, post: dict, delta: dict, output_dir: str) -> list:
    if not HAS_PLOTLY:
        print("  [WARN] plotly not installed — skipping charts.")
        return []
    os.makedirs(output_dir, exist_ok=True)
    charts = []

    nodes      = delta.get("nodes", [])
    node_names = [n["name"] for n in nodes]

    # 1. Node CPU
    if nodes:
        fig = go.Figure()
        fig.add_trace(go.Bar(name="Baseline",    x=node_names, y=[n["cpu_base"] for n in nodes]))
        fig.add_trace(go.Bar(name="Post-Deploy", x=node_names, y=[n["cpu_post"] for n in nodes]))
        fig.update_layout(
            title={"text": "Node CPU: Baseline vs Post-Deploy<br>"
                           "<span style='font-size:14px;font-weight:normal;'>"
                           "CPU cores used per node</span>"},
            barmode="group",
            legend=dict(orientation="h", yanchor="bottom", y=1.05, xanchor="center", x=0.5),
        )
        fig.update_xaxes(title_text="Node")
        fig.update_yaxes(title_text="CPU (cores)")
        path = os.path.join(output_dir, "chart_node_cpu.png")
        fig.write_image(path)
        charts.append(path)

    # 2. Node Memory
    if nodes:
        fig = go.Figure()
        fig.add_trace(go.Bar(name="Baseline",    x=node_names, y=[n["mem_base"] for n in nodes]))
        fig.add_trace(go.Bar(name="Post-Deploy", x=node_names, y=[n["mem_post"] for n in nodes]))
        fig.update_layout(
            title={"text": "Node Memory: Baseline vs Post-Deploy<br>"
                           "<span style='font-size:14px;font-weight:normal;'>"
                           "MiB used per node</span>"},
            barmode="group",
            legend=dict(orientation="h", yanchor="bottom", y=1.05, xanchor="center", x=0.5),
        )
        fig.update_xaxes(title_text="Node")
        fig.update_yaxes(title_text="Memory (MiB)")
        path = os.path.join(output_dir, "chart_node_mem.png")
        fig.write_image(path)
        charts.append(path)

    # 3. Falcon pod breakdown
    falcon_pods = post.get("falcon_pods", [])
    if falcon_pods:
        pod_names = [p["pod"][:28] for p in falcon_pods]
        fig = make_subplots(rows=1, cols=2, subplot_titles=("CPU (cores)", "Memory (MiB)"))
        fig.add_trace(go.Bar(x=pod_names, y=[p["cpu_cores"] for p in falcon_pods], name="CPU"),    row=1, col=1)
        fig.add_trace(go.Bar(x=pod_names, y=[p["mem_mib"]   for p in falcon_pods], name="Memory"), row=1, col=2)
        fig.update_layout(
            title={"text": "Falcon Pod Resource Usage (Post-Deploy)<br>"
                           "<span style='font-size:14px;font-weight:normal;'>"
                           "Per-pod CPU and memory</span>"},
            showlegend=False,
        )
        fig.update_xaxes(tickangle=-30)
        path = os.path.join(output_dir, "chart_falcon_pods.png")
        fig.write_image(path)
        charts.append(path)

    # 4. Webhook latency
    lat = post.get("admission_latency", {})
    lat_vals = {k: v for k, v in lat.items() if k in ("p50", "p95", "p99") and v is not None}
    if lat_vals:
        fig = go.Figure(go.Bar(
            x=list(lat_vals.keys()),
            y=list(lat_vals.values()),
            text=[f"{v} ms" for v in lat_vals.values()],
            textposition="outside",
        ))
        fig.add_hline(y=500,  line_dash="dash", line_color="#ff9900",
                      annotation_text="500ms warn", annotation_position="top right")
        fig.add_hline(y=1000, line_dash="dot",  line_color="#cc0000",
                      annotation_text="1000ms crit", annotation_position="top right")
        fig.update_layout(
            title={"text": "KAC Webhook Latency (Post-Deploy)<br>"
                           "<span style='font-size:14px;font-weight:normal;'>"
                           "p50 / p95 / p99 in milliseconds</span>"},
        )
        fig.update_xaxes(title_text="Percentile")
        fig.update_yaxes(title_text="Latency (ms)")
        path = os.path.join(output_dir, "chart_webhook_latency.png")
        fig.write_image(path)
        charts.append(path)

    return charts


# ─────────────────────────────────────────────────────────────────────────────
# DOCX REPORT
# ─────────────────────────────────────────────────────────────────────────────
def _set_cell_bg(cell, hex_color: str):
    """Set table cell background colour (OOXML shading)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _add_heading(doc: "Document", text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def _add_table(doc: "Document", headers: list, rows: list,
               header_fill: str = "1F3864", header_font: str = "FFFFFF"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(
            int(header_font[0:2], 16),
            int(header_font[2:4], 16),
            int(header_font[4:6], 16),
        )
        _set_cell_bg(hdr_cells[i], header_fill)

    for ridx, row_data in enumerate(rows):
        row_cells = table.add_row().cells
        fill = "EEF0F5" if ridx % 2 == 0 else "FFFFFF"
        for cidx, val in enumerate(row_data):
            row_cells[cidx].text = str(val)
            _set_cell_bg(row_cells[cidx], fill)
    doc.add_paragraph()


def generate_docx_report(
    baseline: dict,
    post: dict,
    delta: dict,
    assessment: dict,
    charts: list,
    output_path: str,
) -> None:
    if not HAS_DOCX:
        print("[ERROR] python-docx not installed — cannot generate DOCX report.")
        return

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    ci = post.get("cluster_info", {})

    # ── Title ─────────────────────────────────────────────────────────────────
    title = doc.add_heading("CrowdStrike Falcon Operator", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("Performance Impact Assessment Report")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    sub.runs[0].bold = True

    doc.add_paragraph()

    # ── Cluster Identity block ────────────────────────────────────────────────
    _add_heading(doc, "1. Cluster Identity", level=1)
    _add_table(
        doc,
        headers=["Field", "Value"],
        rows=[
            ["Cluster Name",        ci.get("cluster_name",  "unknown")],
            ["Kubeconfig Context",  ci.get("context",        "unknown")],
            ["API Server URL",      ci.get("server_url",     "unknown")],
            ["Platform",            ci.get("k8s_platform",   "unknown")],
            ["Kubernetes Version",  ci.get("k8s_version",    "unknown")],
            ["Total Nodes",         str(ci.get("node_count",    "N/A"))],
            ["Worker Nodes",        str(ci.get("worker_count",  "N/A"))],
            ["Report Generated",    datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")],
            ["Baseline Captured",   baseline.get("timestamp", "N/A")],
            ["Post-Deploy Captured", post.get("timestamp",    "N/A")],
        ],
    )

    # ── Risk Assessment ───────────────────────────────────────────────────────
    _add_heading(doc, "2. Risk Assessment", level=1)
    risk = assessment["risk_level"]
    risk_colors = {"LOW": "1A7F37", "MEDIUM": "9A6700", "HIGH": "B91C1C"}
    p = doc.add_paragraph()
    run = p.add_run(f"Overall Risk Level:  {risk}")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(
        int(risk_colors.get(risk, "000000")[0:2], 16),
        int(risk_colors.get(risk, "000000")[2:4], 16),
        int(risk_colors.get(risk, "000000")[4:6], 16),
    )

    for finding in assessment["findings"]:
        bullet = doc.add_paragraph(finding, style="List Bullet")

    doc.add_paragraph()

    # ── Resource Delta Summary ─────────────────────────────────────────────────
    _add_heading(doc, "3. Resource Delta Summary", level=1)
    s = delta.get("summary", {})
    _add_table(
        doc,
        headers=["Metric", "Baseline", "Post-Deploy", "Delta", "% Change"],
        rows=[
            [
                "Cluster CPU (cores)",
                f"{s.get('cluster_cpu_base', 0):.3f}",
                f"{s.get('cluster_cpu_post', 0):.3f}",
                f"{s.get('cluster_cpu_delta', 0):+.3f}",
                f"{s.get('cpu_pct_increase', 0):+.1f}%",
            ],
            [
                "Cluster Memory (MiB)",
                f"{s.get('cluster_mem_base', 0):.0f}",
                f"{s.get('cluster_mem_post', 0):.0f}",
                f"{s.get('cluster_mem_delta_mib', 0):+.0f}",
                f"{s.get('mem_pct_increase', 0):+.1f}%",
            ],
            [
                "Falcon Pods CPU (cores)",
                "N/A",
                f"{s.get('falcon_cpu_cores', 0):.3f}",
                "–", "–",
            ],
            [
                "Falcon Pods Memory (MiB)",
                "N/A",
                f"{s.get('falcon_mem_mib', 0):.0f}",
                "–", "–",
            ],
        ],
    )

    # ── Per-Node Detail ───────────────────────────────────────────────────────
    _add_heading(doc, "4. Per-Node Resource Detail", level=1)
    _add_table(
        doc,
        headers=["Node", "CPU Base", "CPU Post", "CPU Δ", "Mem Base (MiB)", "Mem Post (MiB)", "Mem Δ (MiB)"],
        rows=[
            [
                n["name"],
                f"{n['cpu_base']:.3f}",
                f"{n['cpu_post']:.3f}",
                f"{n['cpu_delta']:+.3f}",
                f"{n['mem_base']:.0f}",
                f"{n['mem_post']:.0f}",
                f"{n['mem_delta_mib']:+.0f}",
            ]
            for n in delta.get("nodes", [])
        ],
    )

    # ── Component Status ──────────────────────────────────────────────────────
    _add_heading(doc, "5. Falcon Component Status", level=1)

    _add_heading(doc, "5.1  FalconNodeSensor (DaemonSet)", level=2)
    ds = post.get("daemonset_status", {})
    if ds:
        _add_table(
            doc,
            headers=["Namespace", "Desired", "Ready", "Available", "Misscheduled", "Update Strategy"],
            rows=[[
                ds.get("namespace", ""),
                str(ds.get("desired", "")),
                str(ds.get("ready", "")),
                str(ds.get("available", "")),
                str(ds.get("misscheduled", "")),
                ds.get("update_strategy", ""),
            ]],
        )
    else:
        doc.add_paragraph("FalconNodeSensor DaemonSet not found in cluster.")

    _add_heading(doc, "5.2  FalconAdmission / KAC (Deployment)", level=2)
    adm = post.get("admission_status", {})
    if adm:
        _add_table(
            doc,
            headers=["Namespace", "Name", "Desired", "Ready", "Available"],
            rows=[[
                adm.get("namespace", ""),
                adm.get("name", ""),
                str(adm.get("desired", "")),
                str(adm.get("ready", "")),
                str(adm.get("available", "")),
            ]],
        )
    else:
        doc.add_paragraph("FalconAdmission deployment not found in cluster.")

    # ── Webhook Config ────────────────────────────────────────────────────────
    _add_heading(doc, "6. Webhook Configuration", level=1)
    wh = post.get("webhook_config", {})
    if wh:
        doc.add_paragraph(f"ValidatingWebhookConfiguration name: {wh.get('name', '')}")
        _add_table(
            doc,
            headers=["Webhook Name", "Failure Policy", "Timeout (s)", "Match Policy", "NS Selector", "Rules"],
            rows=[
                [
                    h["name"],
                    h["failurePolicy"],
                    str(h["timeoutSeconds"]),
                    h["matchPolicy"],
                    "Yes" if h["namespaceSelector"] else "No",
                    str(h["rules"]),
                ]
                for h in wh.get("webhooks", [])
            ],
        )
    else:
        doc.add_paragraph("No Falcon ValidatingWebhookConfiguration found.")

    # ── Admission Latency ─────────────────────────────────────────────────────
    _add_heading(doc, "7. Admission Webhook Latency", level=1)
    lat = post.get("admission_latency", {})
    lat_rows = []
    for k, label in [("p50", "p50 (ms)"), ("p95", "p95 (ms)"), ("p99", "p99 (ms)"), ("mean_ms", "Mean (ms)")]:
        v = lat.get(k)
        if v is not None:
            lat_rows.append([label, str(v)])
    lat_rows.append(["Sample Count", str(lat.get("sample_count", 0))])
    if lat_rows:
        _add_table(doc, headers=["Metric", "Value"], rows=lat_rows)
    else:
        doc.add_paragraph("No latency data available.")

    # ── Falcon Pod Resources ──────────────────────────────────────────────────
    _add_heading(doc, "8. Falcon Pod Resources (Post-Deploy)", level=1)
    falcon_pods = post.get("falcon_pods", [])
    if falcon_pods:
        _add_table(
            doc,
            headers=["Namespace", "Pod", "CPU (cores)", "Memory (MiB)"],
            rows=[
                [p["namespace"], p["pod"], f"{p['cpu_cores']:.3f}", f"{p['mem_mib']:.0f}"]
                for p in falcon_pods
            ],
        )
    else:
        doc.add_paragraph("No Falcon pods found via kubectl top.")

    # ── Warning Events ────────────────────────────────────────────────────────
    events = post.get("events", [])
    if events:
        _add_heading(doc, "9. Warning Events (Falcon Namespaces)", level=1)
        _add_table(
            doc,
            headers=["Namespace", "Reason", "Count", "Last Seen", "Message"],
            rows=[
                [e["namespace"], e["reason"], str(e["count"]), e["last_ts"], e["message"]]
                for e in events
            ],
        )

    # ── Charts ────────────────────────────────────────────────────────────────
    if charts:
        _add_heading(doc, "10. Performance Charts", level=1)
        chart_titles = {
            "chart_node_cpu.png":         "Figure 1 – Node CPU: Baseline vs Post-Deploy",
            "chart_node_mem.png":         "Figure 2 – Node Memory: Baseline vs Post-Deploy",
            "chart_falcon_pods.png":      "Figure 3 – Falcon Pod Resource Breakdown",
            "chart_webhook_latency.png":  "Figure 4 – KAC Admission Webhook Latency",
        }
        for chart_path in charts:
            if os.path.exists(chart_path):
                basename = os.path.basename(chart_path)
                caption  = chart_titles.get(basename, basename)
                doc.add_paragraph(caption).runs[0].bold = True
                doc.add_picture(chart_path, width=Inches(6.0))
                doc.add_paragraph()

    # ── Recommendations ───────────────────────────────────────────────────────
    _add_heading(doc, "11. Recommended Next Steps", level=1)
    recs = [
        "If CPU delta > 10%, review nodeSensorConfig.resources in the FalconNodeSensor CR and lower CPU limits.",
        "If webhook p99 > 500ms, scale FalconAdmission replicas or increase its CPU limits.",
        "Set failurePolicy: Ignore with timeoutSeconds: 5-10 if workload availability is critical.",
        "Add priorityClassName: system-node-critical to FalconNodeSensor pods to prevent eviction.",
        "Re-run this assessment during peak business hours for representative load data.",
        "Monitor container_cpu_cfs_throttled_periods_total to detect sensor CPU throttling.",
        "Enable Prometheus scraping of the falcon-system namespace for continuous trending.",
    ]
    for rec in recs:
        doc.add_paragraph(rec, style="List Number")

    # ── Footer note ───────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_p = doc.add_paragraph(
        f"Report generated by falcon_perf_assessment.py  |  "
        f"Cluster: {ci.get('cluster_name', 'unknown')}  |  "
        f"{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )
    footer_p.runs[0].font.size  = Pt(9)
    footer_p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(output_path)
    print(f"✅  DOCX report written to: {output_path}")


# ─────────────────────────────────────────────────────────────────────────────
# HTML REPORT  (kept from v1, updated with cluster info section)
# ─────────────────────────────────────────────────────────────────────────────
HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>Falcon Operator – Performance Impact Report</title>
<style>
  body { font-family:'Segoe UI',Arial,sans-serif; background:#0f1117; color:#e0e0e0; margin:40px; }
  h1   { color:#e85c27; font-size:2em; }
  h2   { color:#ff7b4f; border-bottom:1px solid #333; padding-bottom:6px; }
  h3   { color:#c0c0c0; }
  table{ width:100%; border-collapse:collapse; margin-bottom:20px; }
  th   { background:#1e2230; color:#ff7b4f; padding:8px 12px; text-align:left; }
  td   { padding:7px 12px; border-bottom:1px solid #2a2d3a; }
  tr:nth-child(even) td { background:#181c29; }
  .badge-low    { background:#1a7f37; color:#fff; padding:2px 10px; border-radius:12px; font-weight:bold; }
  .badge-medium { background:#9a6700; color:#fff; padding:2px 10px; border-radius:12px; font-weight:bold; }
  .badge-high   { background:#b91c1c; color:#fff; padding:2px 10px; border-radius:12px; font-weight:bold; }
  .finding { padding:6px 0; border-bottom:1px dashed #2a2d3a; }
  .chart   { max-width:100%; margin:16px 0; border:1px solid #2a2d3a; border-radius:8px; }
  .meta    { color:#888; font-size:0.85em; }
  .section { background:#13161f; border-radius:8px; padding:16px 24px; margin-bottom:24px; }
  code     { background:#1a1d2b; padding:2px 6px; border-radius:4px; font-family:monospace; }
</style>
</head>
<body>
<h1>🦅 CrowdStrike Falcon Operator – Performance Impact Report</h1>
<p class="meta">Generated: {{ timestamp }}</p>

<div class="section">
  <h2>Cluster Identity</h2>
  <table>
    <tr><th>Field</th><th>Value</th></tr>
    <tr><td>Cluster Name</td><td><strong>{{ ci.cluster_name }}</strong></td></tr>
    <tr><td>Kubeconfig Context</td><td>{{ ci.context }}</td></tr>
    <tr><td>API Server URL</td><td><code>{{ ci.server_url }}</code></td></tr>
    <tr><td>Platform</td><td>{{ ci.k8s_platform }}</td></tr>
    <tr><td>Kubernetes Version</td><td>{{ ci.k8s_version }}</td></tr>
    <tr><td>Total Nodes</td><td>{{ ci.node_count | default('N/A') }}</td></tr>
    <tr><td>Worker Nodes</td><td>{{ ci.worker_count | default('N/A') }}</td></tr>
    <tr><td>Baseline Captured</td><td>{{ baseline_ts }}</td></tr>
    <tr><td>Post-Deploy Captured</td><td>{{ post_ts }}</td></tr>
  </table>
</div>

<div class="section">
  <h2>Risk Assessment</h2>
  <p>Risk Level: <span class="badge-{{ risk_level | lower }}">{{ risk_level }}</span></p>
  {% for f in findings %}<div class="finding">{{ f }}</div>{% endfor %}
</div>

<div class="section">
  <h2>Resource Delta Summary</h2>
  <table>
    <tr><th>Metric</th><th>Baseline</th><th>Post-Deploy</th><th>Delta</th><th>% Change</th></tr>
    <tr>
      <td>Cluster CPU (cores)</td>
      <td>{{ "%.3f"|format(summary.cluster_cpu_base) }}</td>
      <td>{{ "%.3f"|format(summary.cluster_cpu_post) }}</td>
      <td>{{ "%+.3f"|format(summary.cluster_cpu_delta) }}</td>
      <td>{{ "%+.1f"|format(summary.cpu_pct_increase) }}%</td>
    </tr>
    <tr>
      <td>Cluster Memory (MiB)</td>
      <td>{{ "%.0f"|format(summary.cluster_mem_base) }}</td>
      <td>{{ "%.0f"|format(summary.cluster_mem_post) }}</td>
      <td>{{ "%+.0f"|format(summary.cluster_mem_delta_mib) }}</td>
      <td>{{ "%+.1f"|format(summary.mem_pct_increase) }}%</td>
    </tr>
    <tr>
      <td>Falcon Pods CPU (cores)</td><td>N/A</td>
      <td>{{ "%.3f"|format(summary.falcon_cpu_cores) }}</td><td>–</td><td>–</td>
    </tr>
    <tr>
      <td>Falcon Pods Memory (MiB)</td><td>N/A</td>
      <td>{{ "%.0f"|format(summary.falcon_mem_mib) }}</td><td>–</td><td>–</td>
    </tr>
  </table>
</div>

<div class="section">
  <h2>Per-Node Resource Detail</h2>
  <table>
    <tr><th>Node</th><th>CPU Base</th><th>CPU Post</th><th>CPU Δ</th>
        <th>Mem Base (MiB)</th><th>Mem Post (MiB)</th><th>Mem Δ (MiB)</th></tr>
    {% for n in node_delta %}
    <tr>
      <td>{{ n.name }}</td>
      <td>{{ "%.3f"|format(n.cpu_base) }}</td><td>{{ "%.3f"|format(n.cpu_post) }}</td>
      <td>{{ "%+.3f"|format(n.cpu_delta) }}</td>
      <td>{{ "%.0f"|format(n.mem_base) }}</td><td>{{ "%.0f"|format(n.mem_post) }}</td>
      <td>{{ "%+.0f"|format(n.mem_delta_mib) }}</td>
    </tr>
    {% endfor %}
  </table>
</div>

<div class="section">
  <h2>Falcon Component Status</h2>
  <h3>FalconNodeSensor (DaemonSet)</h3>
  {% if ds_status %}
  <table>
    <tr><th>Namespace</th><th>Desired</th><th>Ready</th><th>Available</th><th>Misscheduled</th><th>Update Strategy</th></tr>
    <tr><td>{{ ds_status.namespace }}</td><td>{{ ds_status.desired }}</td>
        <td>{{ ds_status.ready }}</td><td>{{ ds_status.available }}</td>
        <td>{{ ds_status.misscheduled }}</td><td>{{ ds_status.update_strategy }}</td></tr>
  </table>
  {% else %}<p class="meta">FalconNodeSensor not found.</p>{% endif %}

  <h3>FalconAdmission / KAC (Deployment)</h3>
  {% if adm_status %}
  <table>
    <tr><th>Namespace</th><th>Name</th><th>Desired</th><th>Ready</th><th>Available</th></tr>
    <tr><td>{{ adm_status.namespace }}</td><td>{{ adm_status.name }}</td>
        <td>{{ adm_status.desired }}</td><td>{{ adm_status.ready }}</td>
        <td>{{ adm_status.available }}</td></tr>
  </table>
  {% else %}<p class="meta">FalconAdmission not found.</p>{% endif %}
</div>

<div class="section">
  <h2>Webhook Configuration</h2>
  {% if webhook %}
  <p>Name: <code>{{ webhook.name }}</code></p>
  <table>
    <tr><th>Webhook</th><th>Failure Policy</th><th>Timeout (s)</th><th>Match Policy</th><th>NS Selector</th><th>Rules</th></tr>
    {% for wh in webhook.webhooks %}
    <tr><td>{{ wh.name }}</td><td>{{ wh.failurePolicy }}</td><td>{{ wh.timeoutSeconds }}</td>
        <td>{{ wh.matchPolicy }}</td><td>{{ "Yes" if wh.namespaceSelector else "No" }}</td>
        <td>{{ wh.rules }}</td></tr>
    {% endfor %}
  </table>
  {% else %}<p class="meta">No Falcon ValidatingWebhookConfiguration found.</p>{% endif %}
</div>

<div class="section">
  <h2>Admission Webhook Latency</h2>
  {% if lat %}
  <table>
    <tr><th>Metric</th><th>Value</th></tr>
    {% if lat.p50   is not none %}<tr><td>p50 (ms)</td><td>{{ lat.p50 }}</td></tr>{% endif %}
    {% if lat.p95   is not none %}<tr><td>p95 (ms)</td><td>{{ lat.p95 }}</td></tr>{% endif %}
    {% if lat.p99   is not none %}<tr><td>p99 (ms)</td><td>{{ lat.p99 }}</td></tr>{% endif %}
    {% if lat.mean_ms is defined %}<tr><td>Mean (ms)</td><td>{{ lat.mean_ms }}</td></tr>{% endif %}
    <tr><td>Sample Count</td><td>{{ lat.sample_count }}</td></tr>
  </table>
  {% else %}<p class="meta">No latency data available.</p>{% endif %}
</div>

<div class="section">
  <h2>Falcon Pod Resources (Post-Deploy)</h2>
  {% if falcon_pods %}
  <table>
    <tr><th>Namespace</th><th>Pod</th><th>CPU (cores)</th><th>Memory (MiB)</th></tr>
    {% for p in falcon_pods %}
    <tr><td>{{ p.namespace }}</td><td>{{ p.pod }}</td>
        <td>{{ "%.3f"|format(p.cpu_cores) }}</td><td>{{ "%.0f"|format(p.mem_mib) }}</td></tr>
    {% endfor %}
  </table>
  {% else %}<p class="meta">No Falcon pods found via kubectl top.</p>{% endif %}
</div>

{% if events %}
<div class="section">
  <h2>Warning Events</h2>
  <table>
    <tr><th>Namespace</th><th>Reason</th><th>Count</th><th>Last Seen</th><th>Message</th></tr>
    {% for e in events %}
    <tr><td>{{ e.namespace }}</td><td>{{ e.reason }}</td><td>{{ e.count }}</td>
        <td>{{ e.last_ts }}</td><td>{{ e.message }}</td></tr>
    {% endfor %}
  </table>
</div>
{% endif %}

{% if charts %}
<div class="section">
  <h2>Performance Charts</h2>
  {% for c in charts %}<img src="{{ c }}" alt="chart" class="chart"><br>{% endfor %}
</div>
{% endif %}

<div class="section">
  <h2>Recommended Next Steps</h2>
  <ol>
    <li>If CPU delta &gt; 10%, review <code>nodeSensorConfig.resources</code> in the FalconNodeSensor CR.</li>
    <li>If webhook p99 &gt; 500ms, scale FalconAdmission replicas or increase CPU limits.</li>
    <li>Set <code>failurePolicy: Ignore</code> and a short <code>timeoutSeconds</code> (5–10s) for availability.</li>
    <li>Add <code>priorityClassName: system-node-critical</code> on FalconNodeSensor to prevent eviction.</li>
    <li>Re-run during peak load windows for representative data.</li>
    <li>Monitor <code>container_cpu_cfs_throttled_periods_total</code> for sensor throttling.</li>
    <li>Enable Prometheus scraping of falcon-system namespace for continuous trending.</li>
  </ol>
</div>
<p class="meta" style="text-align:center;">
  Report generated by falcon_perf_assessment.py &nbsp;|&nbsp;
  Cluster: {{ ci.cluster_name }} &nbsp;|&nbsp; {{ timestamp }}
</p>
</body></html>"""


def generate_html_report(baseline, post, delta, assessment, charts, output_path):
    if not HAS_JINJA:
        print("[ERROR] jinja2 not installed — cannot render HTML report.")
        return
    report_dir = os.path.dirname(os.path.abspath(output_path))
    rel_charts  = [os.path.relpath(c, report_dir) for c in charts]
    tmpl = Template(HTML_TEMPLATE)
    html = tmpl.render(
        timestamp   = datetime.datetime.now().isoformat(),
        baseline_ts = baseline.get("timestamp", ""),
        post_ts     = post.get("timestamp", ""),
        ci          = post.get("cluster_info", {}),
        risk_level  = assessment["risk_level"],
        findings    = assessment["findings"],
        summary     = delta.get("summary", {}),
        node_delta  = delta.get("nodes", []),
        ds_status   = post.get("daemonset_status", {}),
        adm_status  = post.get("admission_status", {}),
        webhook     = post.get("webhook_config", {}),
        lat         = post.get("admission_latency", {}),
        falcon_pods = post.get("falcon_pods", []),
        events      = post.get("events", []),
        charts      = rel_charts,
    )
    Path(output_path).write_text(html, encoding="utf-8")
    print(f"✅  HTML report written to: {output_path}")


# ─────────────────────────────────────────────────────────────────────────────
# CLI
# ─────────────────────────────────────────────────────────────────────────────
def parse_args():
    p = argparse.ArgumentParser(
        description="Falcon Operator Performance Impact Assessment v2",
        formatter_class=argparse.RawTextHelpFormatter,
        epilog="""
Examples:
  # Capture baseline BEFORE deploying Falcon:
  python3 falcon_perf_assessment.py --mode baseline --output /tmp/baseline.json

  # Generate HTML + DOCX reports AFTER deploying Falcon:
  python3 falcon_perf_assessment.py --mode post \\
      --baseline /tmp/baseline.json \\
      --output /tmp/falcon_report

  # Continuous sampling (optional):
  python3 falcon_perf_assessment.py --mode sample \\
      --samples 12 --interval 60 --output /tmp/samples.json
        """,
    )
    p.add_argument("--mode",     required=True, choices=["baseline", "post", "sample"])
    p.add_argument("--baseline", default=None,
                   help="Path to baseline JSON (required for --mode post)")
    p.add_argument("--output",   required=True,
                   help=("Output path. For 'post' mode omit extension — "
                         "script writes both .html and .docx automatically."))
    p.add_argument("--samples",   type=int, default=6)
    p.add_argument("--interval",  type=int, default=60)
    p.add_argument("--chart-dir", default=None)
    return p.parse_args()


def main():
    args = parse_args()

    if args.mode == "baseline":
        snap = collect_snapshot(label="baseline")
        Path(args.output).write_text(json.dumps(snap, indent=2))
        print(f"\n✅  Baseline saved to: {args.output}")

    elif args.mode == "sample":
        samples = collect_samples(args.samples, args.interval)
        Path(args.output).write_text(json.dumps(samples, indent=2))
        print(f"\n✅  {len(samples)} samples saved to: {args.output}")

    elif args.mode == "post":
        if not args.baseline:
            print("[ERROR] --baseline <path> is required for --mode post")
            sys.exit(1)
        baseline_path = Path(args.baseline)
        if not baseline_path.exists():
            print(f"[ERROR] Baseline file not found: {args.baseline}")
            sys.exit(1)

        baseline = json.loads(baseline_path.read_text())
        post     = collect_snapshot(label="post")
        delta    = compute_delta(baseline, post)
        assess   = assess_risk(delta, post)

        # Strip extension from --output so we can add .html / .docx
        base_out  = str(args.output).removesuffix(".html").removesuffix(".docx")
        chart_dir = args.chart_dir or str(Path(base_out).parent)
        charts    = generate_charts(baseline, post, delta, chart_dir)

        generate_html_report(baseline, post, delta, assess, charts, base_out + ".html")
        generate_docx_report(baseline, post, delta, assess, charts, base_out + ".docx")

        print("\n── Assessment Summary ─────────────────────────────────────")
        ci = post.get("cluster_info", {})
        print(f"  Cluster    : {ci.get('cluster_name')}  ({ci.get('k8s_platform')})")
        print(f"  Version    : {ci.get('k8s_version')}")
        print(f"  Risk Level : {assess['risk_level']}")
        for f in assess["findings"]:
            print(f"  {f}")
        print(f"  HTML  → {base_out}.html")
        print(f"  DOCX  → {base_out}.docx")
        print("────────────────────────────────────────────────────────────")


if __name__ == "__main__":
    main()
