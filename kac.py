#!/usr/bin/env python3
"""
falcon_kac_perf_assessment.py
Performance Impact Assessment – CrowdStrike Falcon KAC (Helm deployment)

Usage:
  # Step 1 – BEFORE helm install (no falcon-kac needed in cluster)
  python3 falcon_kac_perf_assessment.py --mode baseline --output output/baseline.json

  # Step 2 – Install KAC via Helm, then:
  python3 falcon_kac_perf_assessment.py --mode post \
      --baseline output/baseline.json \
      --output output/kac_report \
      --namespace falcon-kac \
      --release falcon-kac

  # Optional continuous sampling
  python3 falcon_kac_perf_assessment.py --mode sample \
      --samples 10 --interval 30 --output output/samples.json

Dependencies:
  pip install jinja2 plotly kaleido python-docx
"""

import argparse
import json
import subprocess
import sys
import time
import datetime
import os
from pathlib import Path

try:
    import plotly.graph_objects as go
    from plotly.subplots import make_subplots
    HAS_PLOTLY = True
except ImportError:
    HAS_PLOTLY = False

try:
    from jinja2 import Template
    HAS_JINJA = True
except ImportError:
    HAS_JINJA = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ─────────────────────────────────────────────────────────────────────────────
# CONSTANTS
# ─────────────────────────────────────────────────────────────────────────────
DEFAULT_NAMESPACE = "falcon-system"
DEFAULT_RELEASE   = "falcon-kac"
DEPLOY_NAME       = "falcon-kac"   # deployment name inside falcon-system namespace
WEBHOOK_PATTERN   = "falcon"


# ─────────────────────────────────────────────────────────────────────────────
# KUBECTL HELPERS
# ─────────────────────────────────────────────────────────────────────────────
def run_kubectl(args, timeout=30):
    try:
        r = subprocess.run(["kubectl"] + args, capture_output=True, text=True, timeout=timeout)
        return r.stdout, r.stderr, r.returncode
    except subprocess.TimeoutExpired:
        return "", f"kubectl timeout after {timeout}s", 1
    except FileNotFoundError:
        return "", "kubectl not found in PATH", 1


def kubectl_json(args):
    stdout, stderr, rc = run_kubectl(args + ["-o", "json"])
    if rc != 0:
        print(f"  [WARN] {stderr.strip()}", file=sys.stderr)
        return None
    try:
        return json.loads(stdout)
    except json.JSONDecodeError:
        return None


def kubectl_raw(path):
    stdout, _, _ = run_kubectl(["get", "--raw", path])
    return stdout


# ─────────────────────────────────────────────────────────────────────────────
# PARSING HELPERS
# ─────────────────────────────────────────────────────────────────────────────
def _parse_cpu(raw):
    raw = raw.strip()
    if raw.endswith("m"):
        return float(raw[:-1]) / 1000
    try:
        return float(raw)
    except ValueError:
        return 0.0


def _parse_mem(raw):
    raw = raw.strip()
    if   raw.endswith("Ki"): return float(raw[:-2]) / 1024
    elif raw.endswith("Mi"): return float(raw[:-2])
    elif raw.endswith("Gi"): return float(raw[:-2]) * 1024
    elif raw.endswith("Ti"): return float(raw[:-2]) * 1024 * 1024
    try:
        return float(raw) / (1024 * 1024)
    except ValueError:
        return 0.0


def _safe_diff(a, b):
    if a is None or b is None:
        return None
    return round(a - b, 2)



# ─────────────────────────────────────────────────────────────────────────────
# FLUXCD HELPERS  (primary source of truth when deployed via HelmRelease CRD)
# ─────────────────────────────────────────────────────────────────────────────
FLUX_NS = "flux-system"   # namespace where Flux controllers + HelmChart CRDs live

def collect_flux_helmrelease(namespace, release):
    """
    Read FluxCD HelmRelease CR — source of truth for version + status
    when the chart was installed via Flux GitRepository → HelmChart → HelmRelease.

    kubectl get helmrelease <release> -n <namespace> -o json
    """
    hr = kubectl_json(["get", "helmrelease", release, "-n", namespace])
    if not hr:
        return {}

    status   = hr.get("status",  {})
    spec     = hr.get("spec",    {})
    history  = status.get("history", [])
    latest   = history[0] if history else {}

    # Conditions
    conditions = {
        c["type"]: {"status": c["status"], "reason": c.get("reason",""), "message": c.get("message","")}
        for c in status.get("conditions", [])
    }
    ready_cond   = conditions.get("Ready",    {})
    released_cond = conditions.get("Released",{})

    # Derive overall status: prefer Flux condition over helm secret
    if ready_cond.get("status") == "True":
        flux_status = "deployed"
    elif ready_cond.get("status") == "False":
        flux_status = f"flux-failed: {ready_cond.get('reason','')}"
    else:
        flux_status = "unknown"

    chart_spec    = spec.get("chart", {}).get("spec", {})
    chart_ref     = chart_spec.get("chart", "unknown")     # e.g. charts/falcon-kac/v1.6.0
    chart_version = chart_spec.get("version", "*")         # semver constraint or *
    source_ref    = chart_spec.get("sourceRef", {})

    return {
        "source":          "FluxCD HelmRelease",
        "release":         release,
        "namespace":       namespace,
        # Version from last successful history entry (most reliable)
        "chart":           latest.get("chartName",    chart_ref),
        "chart_version":   latest.get("chartVersion", chart_version),
        "app_version":     latest.get("appVersion",   "unknown"),
        "status":          latest.get("status",       flux_status),
        "deployed_at":     latest.get("lastDeployed", status.get("lastAppliedRevision", "unknown")),
        "flux_status":     flux_status,
        "flux_ready":      ready_cond.get("status") == "True",
        "flux_reason":     ready_cond.get("reason",  ""),
        "flux_message":    ready_cond.get("message", ""),
        "released_status": released_cond.get("status",""),
        "released_reason": released_cond.get("reason",""),
        "git_source":      source_ref.get("name",""),
        "chart_path":      chart_ref,
        "chart_selector":  chart_version,
        "reconcile_at":    status.get("lastHandledReconcileAt",""),
        "observed_gen":    status.get("observedGeneration",0),
        "history":         history[:3],   # keep last 3 entries
        "conditions":      conditions,
        "values_summary":  {},            # populated below via helm get values fallback
    }


def collect_flux_helmchart(release, flux_ns=FLUX_NS):
    """
    Read FluxCD HelmChart CR in flux-system.
    Name convention: <namespace>-<release>  e.g. falcon-system-falcon-kac
    """
    # Try common naming patterns
    for candidate in [
        f"{DEFAULT_NAMESPACE}-{release}",
        f"falcon-system-{release}",
        f"falcon-kac-helm-chart",
        release,
    ]:
        hc = kubectl_json(["get", "helmchart", candidate, "-n", flux_ns])
        if hc:
            status = hc.get("status", {})
            spec   = hc.get("spec",   {})
            return {
                "name":            candidate,
                "chart":           spec.get("chart",  ""),
                "version":         spec.get("version","*"),
                "source_kind":     spec.get("sourceRef",{}).get("kind",""),
                "source_name":     spec.get("sourceRef",{}).get("name",""),
                "observed_version":status.get("observedChartName",""),
                "artifact_path":   status.get("artifact",{}).get("path",""),
                "revision":        status.get("artifact",{}).get("revision",""),
                "ready":           any(
                    c.get("type") == "Ready" and c.get("status") == "True"
                    for c in status.get("conditions",[])
                ),
                "conditions": [
                    {"type": c["type"], "status": c["status"], "reason": c.get("reason","")}
                    for c in status.get("conditions",[])
                ],
            }
    return {}


def _enrich_helm_values(info, namespace, release):
    """Attempt helm get values to add values_summary to an existing info dict."""
    try:
        r = subprocess.run(
            ["helm", "get", "values", release, "-n", namespace, "--output", "json"],
            capture_output=True, text=True, timeout=20
        )
        if r.returncode == 0:
            vals = json.loads(r.stdout)
            info["values_summary"] = {
                "replicas":          vals.get("replicas", "default"),
                "resources_ac":      vals.get("resources", {}).get("falcon-ac", {}),
                "resources_client":  vals.get("resources", {}).get("falcon-client", {}),
                "failurePolicy":     vals.get("webhook", {}).get("failurePolicy", "not set"),
                "timeoutSeconds":    vals.get("webhook", {}).get("timeoutSeconds", "not set"),
                "priorityClassName": vals.get("priorityClassName", "not set"),
            }
    except (FileNotFoundError, subprocess.TimeoutExpired, json.JSONDecodeError):
        pass
    return info

# ─────────────────────────────────────────────────────────────────────────────
# CLUSTER IDENTITY
# ─────────────────────────────────────────────────────────────────────────────
def collect_cluster_info():
    info = {
        "cluster_name": "unknown", "context": "unknown",
        "server_url": "unknown",   "k8s_version": "unknown",
        "k8s_platform": "unknown", "node_count": "N/A",
        "worker_count": "N/A",
    }
    stdout, _, rc = run_kubectl(["config", "current-context"])
    if rc == 0:
        info["context"] = stdout.strip()

    cfg = kubectl_json(["config", "view"])
    if cfg:
        for ctx in cfg.get("contexts", []):
            if ctx.get("name") == info["context"]:
                info["cluster_name"] = ctx.get("context", {}).get("cluster", info["context"])
                break
        for cl in cfg.get("clusters", []):
            if cl.get("name") == info["cluster_name"]:
                info["server_url"] = cl.get("cluster", {}).get("server", "unknown")
                break

    ver = kubectl_json(["version"])
    if ver:
        sv = ver.get("serverVersion", {})
        info["k8s_version"] = sv.get("gitVersion", "unknown")
        gv = sv.get("gitVersion", "").lower()
        if   "eks"       in gv: info["k8s_platform"] = "Amazon EKS"
        elif "gke"       in gv: info["k8s_platform"] = "Google GKE"
        elif "aks"       in gv: info["k8s_platform"] = "Azure AKS"
        elif "openshift" in gv: info["k8s_platform"] = "OpenShift"
        else:
            nodes_obj = kubectl_json(["get", "nodes"])
            if nodes_obj:
                for n in nodes_obj.get("items", []):
                    pid = n.get("spec", {}).get("providerID", "")
                    if   "eks"   in pid: info["k8s_platform"] = "Amazon EKS";  break
                    elif "gce"   in pid: info["k8s_platform"] = "Google GKE";  break
                    elif "azure" in pid: info["k8s_platform"] = "Azure AKS";   break
                    elif n.get("metadata",{}).get("labels",{}).get("node.openshift.io/os_id"):
                        info["k8s_platform"] = "OpenShift"; break
                else:
                    info["k8s_platform"] = "Vanilla / On-Prem"

    nodes_obj = kubectl_json(["get", "nodes"])
    if nodes_obj:
        items = nodes_obj.get("items", [])
        info["node_count"] = len(items)
        info["worker_count"] = sum(
            1 for n in items
            if not n.get("metadata",{}).get("labels",{}).get("node-role.kubernetes.io/control-plane")
            and not n.get("metadata",{}).get("labels",{}).get("node-role.kubernetes.io/master")
        )
    return info


# ─────────────────────────────────────────────────────────────────────────────
# CLUSTER-WIDE METRICS (used in both baseline and post)
# ─────────────────────────────────────────────────────────────────────────────
def collect_node_resources():
    stdout, stderr, rc = run_kubectl(["top", "nodes", "--no-headers"])
    nodes = []
    if rc != 0:
        print(f"  [WARN] top nodes failed: {stderr.strip()}", file=sys.stderr)
        return nodes
    for line in stdout.strip().splitlines():
        parts = line.split()
        if len(parts) < 5:
            continue
        nodes.append({
            "name":      parts[0],
            "cpu_cores": _parse_cpu(parts[1]),
            "mem_mib":   _parse_mem(parts[3]),
        })
    return nodes


def collect_pod_resources(namespace=""):
    args = ["top", "pods", "--no-headers"]
    args += ["-n", namespace] if namespace else ["-A"]
    stdout, stderr, rc = run_kubectl(args)
    pods = []
    if rc != 0:
        print(f"  [WARN] top pods failed: {stderr.strip()}", file=sys.stderr)
        return pods
    for line in stdout.strip().splitlines():
        parts = line.split()
        if len(parts) < 4:
            continue
        if namespace:
            ns, pod, cpu_raw, mem_raw = namespace, parts[0], parts[1], parts[2]
        else:
            ns, pod, cpu_raw, mem_raw = parts[0], parts[1], parts[2], parts[3]
        pods.append({
            "namespace": ns, "pod": pod,
            "cpu_cores": _parse_cpu(cpu_raw),
            "mem_mib":   _parse_mem(mem_raw),
        })
    return pods


def collect_node_capacity():
    nodes_obj = kubectl_json(["get", "nodes"])
    if not nodes_obj:
        return []
    result = []
    for item in nodes_obj.get("items", []):
        name        = item["metadata"]["name"]
        allocatable = item.get("status", {}).get("allocatable", {})
        capacity    = item.get("status", {}).get("capacity", {})
        result.append({
            "name":                name,
            "allocatable_cpu":     _parse_cpu(allocatable.get("cpu", "0")),
            "allocatable_mem_mib": _parse_mem(allocatable.get("memory", "0Ki")),
            "capacity_cpu":        _parse_cpu(capacity.get("cpu", "0")),
            "capacity_mem_mib":    _parse_mem(capacity.get("memory", "0Ki")),
        })
    return result


def collect_apiserver_latency():
    raw = kubectl_raw("/metrics")
    result = {"post_p99_bound_ms": None, "create_p99_bound_ms": None}
    if not raw:
        return result
    for line in raw.splitlines():
        if "apiserver_request_duration_seconds_bucket" not in line or line.startswith("#"):
            continue
        for verb in ("POST", "CREATE"):
            if f'verb="{verb}"' not in line:
                continue
            try:
                le_s = line.index('le="') + 4
                le_e = line.index('"', le_s)
                le_v = line[le_s:le_e]
                if le_v != "+Inf":
                    result[f"{verb.lower()}_p99_bound_ms"] = round(float(le_v) * 1000, 2)
            except (ValueError, IndexError):
                pass
    return result


# ─────────────────────────────────────────────────────────────────────────────
# HELM RELEASE INFO  (FluxCD-aware)
# ─────────────────────────────────────────────────────────────────────────────
def collect_helm_release(namespace, release):
    """
    Primary: query FluxCD HelmRelease CRD (accurate for Flux-managed installs).
    Fallback: helm list → helm status (for plain helm install).
    """
    # ── Try FluxCD HelmRelease first ────────────────────────────────────────
    flux_info = collect_flux_helmrelease(namespace, release)
    if flux_info:
        print(f"     [Flux] chart={flux_info['chart_version']}  "
              f"appVersion={flux_info['app_version']}  "
              f"status={flux_info['flux_status']}")
        return _enrich_helm_values(flux_info, namespace, release)

    # ── Fallback: helm list (reliable even when helm status shows failed) ────
    info = {
        "source": "helm list", "release": release, "namespace": namespace,
        "chart": "unknown", "chart_version": "unknown",
        "app_version": "unknown", "status": "unknown",
        "deployed_at": "unknown", "values_summary": {},
        "flux_status": "n/a", "flux_ready": False,
        "flux_reason": "", "flux_message": "",
        "git_source": "", "chart_path": "", "chart_selector": "",
    }
    try:
        r = subprocess.run(
            ["helm", "list", "-n", namespace, "--filter", f"^{release}$", "--output", "json"],
            capture_output=True, text=True, timeout=20
        )
        if r.returncode == 0:
            items = json.loads(r.stdout)
            if items:
                raw_chart = items[0].get("chart", "unknown")
                # chart field is "name-version" e.g. "falcon-kac-1.6.0"
                parts = raw_chart.rsplit("-", 1)
                info["chart"]         = parts[0] if len(parts) == 2 else raw_chart
                info["chart_version"] = parts[1] if len(parts) == 2 else "unknown"
                info["app_version"]   = items[0].get("app_version", "unknown")
                info["status"]        = items[0].get("status",      "unknown")
                info["deployed_at"]   = items[0].get("updated",     "unknown")
                print(f"     [helm list] chart={info['chart_version']}  "
                      f"appVersion={info['app_version']}  status={info['status']}")
                return _enrich_helm_values(info, namespace, release)
    except (FileNotFoundError, subprocess.TimeoutExpired, json.JSONDecodeError):
        pass

    # ── Last resort: helm status JSON ────────────────────────────────────────
    try:
        r = subprocess.run(
            ["helm", "status", release, "-n", namespace, "--output", "json"],
            capture_output=True, text=True, timeout=20
        )
        if r.returncode == 0:
            hs = json.loads(r.stdout)
            meta = hs.get("chart", {}).get("metadata", {})
            info["chart"]         = meta.get("name",       "unknown")
            info["chart_version"] = meta.get("version",    "unknown")
            info["app_version"]   = meta.get("appVersion", "unknown")
            info["status"]        = hs.get("info", {}).get("status", "unknown")
            info["deployed_at"]   = hs.get("info", {}).get("last_deployed", "unknown")
            print(f"     [helm status] chart={info['chart_version']}  "
                  f"appVersion={info['app_version']}  status={info['status']}")
    except (FileNotFoundError, subprocess.TimeoutExpired, json.JSONDecodeError):
        pass

    return _enrich_helm_values(info, namespace, release)

# ─────────────────────────────────────────────────────────────────────────────
# KAC-SPECIFIC METRIC COLLECTION
# ─────────────────────────────────────────────────────────────────────────────
def collect_kac_deployment(namespace):
    dep = kubectl_json(["get", "deployment", DEPLOY_NAME, "-n", namespace])
    if not dep:
        return {}
    status = dep.get("status", {})
    spec   = dep.get("spec",   {})
    containers = spec.get("template", {}).get("spec", {}).get("containers", [])
    container_resources = {}
    for c in containers:
        name = c.get("name", "unknown")
        res  = c.get("resources", {})
        container_resources[name] = {
            "requests": res.get("requests", {}),
            "limits":   res.get("limits",   {}),
        }
    return {
        "namespace":           namespace,
        "name":                DEPLOY_NAME,
        "desired":             spec.get("replicas", 1),
        "ready":               status.get("readyReplicas", 0),
        "available":           status.get("availableReplicas", 0),
        "updated":             status.get("updatedReplicas", 0),
        "unavailable":         status.get("unavailableReplicas", 0),
        "container_resources": container_resources,
        "conditions": [
            {"type": c["type"], "status": c["status"], "reason": c.get("reason", "")}
            for c in status.get("conditions", [])
        ],
    }


def collect_kac_pods(namespace):
    pods = []
    stdout, _, rc = run_kubectl(["top", "pods", "-n", namespace, "--no-headers"])
    pod_totals = {}
    if rc == 0:
        for line in stdout.strip().splitlines():
            parts = line.split()
            if len(parts) >= 3:
                pod_totals[parts[0]] = {
                    "cpu_cores": _parse_cpu(parts[1]),
                    "mem_mib":   _parse_mem(parts[2]),
                }

    stdout, _, rc = run_kubectl(
        ["top", "pods", "-n", namespace, "--containers", "--no-headers"]
    )
    pod_containers = {}
    if rc == 0:
        for line in stdout.strip().splitlines():
            parts = line.split()
            if len(parts) >= 4:
                pname, cname, cpu_raw, mem_raw = parts[0], parts[1], parts[2], parts[3]
                pod_containers.setdefault(pname, []).append({
                    "container": cname,
                    "cpu_cores": _parse_cpu(cpu_raw),
                    "mem_mib":   _parse_mem(mem_raw),
                })

    pods_obj = kubectl_json(["get", "pods", "-n", namespace])
    pod_meta = {}
    if pods_obj:
        for p in pods_obj.get("items", []):
            pname = p["metadata"]["name"]
            cs    = p.get("status", {}).get("containerStatuses", [])
            pod_meta[pname] = {
                "node":     p.get("spec", {}).get("nodeName", "unknown"),
                "phase":    p.get("status", {}).get("phase", "unknown"),
                "restarts": sum(c.get("restartCount", 0) for c in cs),
            }

    all_pod_names = set(list(pod_totals.keys()) + list(pod_meta.keys()))
    for pname in all_pod_names:
        if DEPLOY_NAME not in pname:
            continue
        totals = pod_totals.get(pname, {"cpu_cores": 0, "mem_mib": 0})
        meta   = pod_meta.get(pname, {})
        pods.append({
            "pod":        pname,
            "namespace":  namespace,
            "node":       meta.get("node", "unknown"),
            "phase":      meta.get("phase", "unknown"),
            "restarts":   meta.get("restarts", 0),
            "cpu_cores":  totals["cpu_cores"],
            "mem_mib":    totals["mem_mib"],
            "containers": pod_containers.get(pname, []),
        })
    return pods


def collect_webhook_config():
    vwcs = kubectl_json(["get", "validatingwebhookconfigurations"])
    if not vwcs:
        return {}
    for item in vwcs.get("items", []):
        name = item.get("metadata", {}).get("name", "")
        if WEBHOOK_PATTERN in name.lower():
            return {
                "name": name,
                "webhooks": [
                    {
                        "name":             wh.get("name", ""),
                        "failurePolicy":    wh.get("failurePolicy", "unknown"),
                        "timeoutSeconds":   wh.get("timeoutSeconds", "default"),
                        "matchPolicy":      wh.get("matchPolicy", "Equivalent"),
                        "sideEffects":      wh.get("sideEffects", "None"),
                        "namespaceSelector": bool(wh.get("namespaceSelector")),
                        "rules":            len(wh.get("rules", [])),
                    }
                    for wh in item.get("webhooks", [])
                ],
            }
    return {}


def collect_kac_admission_latency():
    raw = kubectl_raw("/metrics")
    lat = {"p50": None, "p95": None, "p99": None, "sample_count": 0, "mean_ms": None}
    if not raw:
        return lat
    buckets, total_count, total_sum = {}, 0, 0.0
    for line in raw.splitlines():
        if "apiserver_admission_webhook_admission_duration_seconds" not in line:
            continue
        if WEBHOOK_PATTERN not in line.lower() or line.startswith("#"):
            continue
        if "_bucket{" in line:
            try:
                le_s = line.index('le="') + 4
                le_e = line.index('"', le_s)
                le_v = line[le_s:le_e]
                if le_v != "+Inf":
                    buckets[float(le_v)] = float(line.split()[-1])
            except (ValueError, IndexError):
                pass
        elif "_count{" in line:
            try:   total_count = int(float(line.split()[-1]))
            except (ValueError, IndexError): pass
        elif "_sum{" in line:
            try:   total_sum = float(line.split()[-1])
            except (ValueError, IndexError): pass
    if total_count > 0:
        lat["sample_count"] = total_count
        lat["mean_ms"] = round(total_sum / total_count * 1000, 2)
    if buckets and total_count > 0:
        for pct, key in [(0.5, "p50"), (0.95, "p95"), (0.99, "p99")]:
            target = pct * total_count
            for bound in sorted(buckets):
                if buckets[bound] >= target:
                    lat[key] = round(bound * 1000, 2)
                    break
    return lat


def collect_kac_hpa(namespace):
    hpa = kubectl_json(["get", "hpa", DEPLOY_NAME, "-n", namespace])
    if not hpa:
        return {}
    spec, status = hpa.get("spec", {}), hpa.get("status", {})
    return {
        "min_replicas": spec.get("minReplicas", 1),
        "max_replicas": spec.get("maxReplicas", 1),
        "current":      status.get("currentReplicas", 0),
        "desired":      status.get("desiredReplicas", 0),
        "metrics":      [m.get("type") for m in spec.get("metrics", [])],
    }


def collect_resource_quota(namespace):
    rqs = kubectl_json(["get", "resourcequota", "-n", namespace])
    if not rqs or not rqs.get("items"):
        return {}
    out = {}
    for rq in rqs.get("items", []):
        name   = rq["metadata"]["name"]
        status = rq.get("status", {})
        out[name] = {"hard": status.get("hard", {}), "used": status.get("used", {})}
    return out


def collect_events(namespace):
    evts = kubectl_json(["get", "events", "-n", namespace, "--sort-by=.lastTimestamp"])
    if not evts:
        return []
    out = []
    for e in evts.get("items", [])[-20:]:
        out.append({
            "type":    e.get("type", "Normal"),
            "reason":  e.get("reason", ""),
            "message": e.get("message", "")[:150],
            "count":   e.get("count", 1),
            "last_ts": e.get("lastTimestamp", ""),
            "object":  e.get("involvedObject", {}).get("name", ""),
        })
    warnings = [e for e in out if e["type"] == "Warning"]
    return warnings if warnings else out[-5:]


def _wait_for_kac_ready(namespace, release=DEFAULT_RELEASE, timeout_s=120):
    """
    Wait for KAC to be ready via two parallel checks:
    1. kubectl rollout status (deployment readiness)
    2. FluxCD HelmRelease Ready condition (Flux reconciliation success)
    """
    deadline = time.time() + timeout_s

    # First, quickly check if the Flux HelmRelease is already marked Ready
    hr = kubectl_json(["get", "helmrelease", release, "-n", namespace])
    if hr:
        for cond in hr.get("status", {}).get("conditions", []):
            if cond.get("type") == "Ready" and cond.get("status") == "True":
                print("     Flux HelmRelease is Ready — skipping rollout wait.")
                return

    while time.time() < deadline:
        # Check rollout
        stdout, _, rc = run_kubectl(
            ["rollout", "status", "deployment", DEPLOY_NAME,
             "-n", namespace, "--timeout=10s"]
        )
        if rc == 0 and "successfully rolled out" in stdout.lower():
            print("     KAC deployment is ready.")
            return

        # Also check Flux HelmRelease condition mid-wait
        hr = kubectl_json(["get", "helmrelease", release, "-n", namespace])
        if hr:
            for cond in hr.get("status", {}).get("conditions", []):
                if cond.get("type") == "Ready" and cond.get("status") == "True":
                    print("     Flux HelmRelease became Ready.")
                    return

        remaining = int(deadline - time.time())
        print(f"     Waiting for KAC rollout... ({remaining}s remaining)")
        time.sleep(10)
    print(f"  [WARN] KAC not ready after {timeout_s}s — metrics may be incomplete.")


# ─────────────────────────────────────────────────────────────────────────────
# SNAPSHOTS
# ─────────────────────────────────────────────────────────────────────────────
def collect_baseline_snapshot(namespace, release):
    """Cluster-wide only — runs BEFORE falcon-kac is installed."""
    print(f"\n[BASELINE] {datetime.datetime.now().isoformat()}")
    print("  NOTE: Collecting cluster-wide metrics only (falcon-kac not yet installed).")
    snap = _empty_snap("baseline", namespace, release)

    print("  → Cluster identity ...")
    snap["cluster_info"] = collect_cluster_info()
    ci = snap["cluster_info"]
    print(f"     Cluster : {ci['cluster_name']}  |  Platform: {ci['k8s_platform']}  |  K8s: {ci['k8s_version']}")

    print("  → Node resource usage ...")
    snap["node_resources"] = collect_node_resources()

    print("  → Node capacity / allocatable ...")
    snap["node_capacity"] = collect_node_capacity()

    print("  → All pod resource usage (cluster-wide) ...")
    snap["all_pods"] = collect_pod_resources()

    print("  → API server latency baseline ...")
    snap["apiserver_latency"] = collect_apiserver_latency()

    print(f"\n  Baseline captured: {len(snap['node_resources'])} nodes, {len(snap['all_pods'])} pods total.")
    print("  Next: helm upgrade --install falcon-kac ... then run --mode post\n")
    return snap


def collect_post_snapshot(namespace, release):
    """Full collection — runs AFTER falcon-kac is installed."""
    print(f"\n[POST-DEPLOY] {datetime.datetime.now().isoformat()}")
    snap = _empty_snap("post", namespace, release)

    print("  → Cluster identity ...")
    snap["cluster_info"] = collect_cluster_info()
    ci = snap["cluster_info"]
    print(f"     Cluster : {ci['cluster_name']}  |  Platform: {ci['k8s_platform']}  |  K8s: {ci['k8s_version']}")

    print(f"  → Helm release [{release}] in namespace [{namespace}] ...")
    snap["helm_release"] = collect_helm_release(namespace, release)
    hr = snap["helm_release"]
    if hr.get("status") not in ("deployed", "unknown"):
        print(f"  [WARN] Helm status is '{hr.get('status')}' — KAC may not be fully deployed.")
    print(f"     Chart: {hr['chart']}-{hr['chart_version']}  |  Status: {hr['status']}  |  AppVersion: {hr['app_version']}")

    print("  → Waiting for KAC deployment to be available ...")
    _wait_for_kac_ready(namespace, release)

    print("  → FluxCD HelmChart status (flux-system) ...")
    snap["flux_helmchart"] = collect_flux_helmchart(release)

    print("  → KAC Deployment status ...")
    snap["kac_deployment"] = collect_kac_deployment(namespace)

    print("  → KAC pod resources (kubectl top pods --containers) ...")
    snap["kac_pods"] = collect_kac_pods(namespace)

    print("  → ValidatingWebhookConfiguration ...")
    snap["webhook_config"] = collect_webhook_config()

    print("  → Admission webhook latency from apiserver /metrics ...")
    snap["admission_latency"] = collect_kac_admission_latency()

    print("  → Node resource usage (post-deploy) ...")
    snap["node_resources"] = collect_node_resources()

    print("  → Node capacity / allocatable ...")
    snap["node_capacity"] = collect_node_capacity()

    print("  → All pod resource usage (post-deploy) ...")
    snap["all_pods"] = collect_pod_resources()

    print("  → API server latency (post-deploy) ...")
    snap["apiserver_latency"] = collect_apiserver_latency()

    print("  → HPA (if configured) ...")
    snap["hpa"] = collect_kac_hpa(namespace)

    print("  → ResourceQuota ...")
    snap["resource_quota"] = collect_resource_quota(namespace)

    print("  → Events ...")
    snap["events"] = collect_events(namespace)

    return snap


def _empty_snap(label, namespace, release):
    return {
        "label": label, "timestamp": datetime.datetime.now().isoformat(),
        "namespace": namespace, "release": release,
        "cluster_info": {}, "helm_release": {},
        "flux_helmchart": {}, "kac_deployment": {}, "kac_pods": [],
        "webhook_config": {}, "admission_latency": {},
        "hpa": {}, "resource_quota": {}, "events": [],
        "node_resources": [], "node_capacity": [],
        "all_pods": [], "apiserver_latency": {},
    }


def collect_samples(n, interval_s, namespace, release):
    samples = []
    for i in range(n):
        print(f"\n── Sample {i+1}/{n} ────────────────────────────────────────")
        samples.append(collect_post_snapshot(namespace, release))
        if i < n - 1:
            print(f"  Sleeping {interval_s}s ...")
            time.sleep(interval_s)
    return samples


# ─────────────────────────────────────────────────────────────────────────────
# DELTA
# ─────────────────────────────────────────────────────────────────────────────
def compute_delta(baseline, post):
    delta = {"nodes": [], "pods": [], "summary": {}}

    # Node-level delta (cluster overhead)
    base_nodes = {n["name"]: n for n in baseline.get("node_resources", [])}
    post_nodes = {n["name"]: n for n in post.get("node_resources", [])}
    for name, pn in post_nodes.items():
        bn = base_nodes.get(name, {"cpu_cores": 0, "mem_mib": 0})
        delta["nodes"].append({
            "name":          name,
            "cpu_base":      bn["cpu_cores"],
            "cpu_post":      pn["cpu_cores"],
            "cpu_delta":     round(pn["cpu_cores"] - bn["cpu_cores"], 4),
            "mem_base":      bn["mem_mib"],
            "mem_post":      pn["mem_mib"],
            "mem_delta_mib": round(pn["mem_mib"] - bn["mem_mib"], 2),
        })

    # KAC pod delta (post only — not present at baseline)
    for p in post.get("kac_pods", []):
        delta["pods"].append({
            "pod":       p["pod"],
            "node":      p.get("node", ""),
            "cpu_base":  0,
            "cpu_post":  p["cpu_cores"],
            "cpu_delta": p["cpu_cores"],
            "mem_base":  0,
            "mem_post":  p["mem_mib"],
            "mem_delta": p["mem_mib"],
            "containers": p.get("containers", []),
        })

    # Cluster totals
    total_base_cpu = sum(n["cpu_cores"] for n in baseline.get("node_resources", []))
    total_post_cpu = sum(n["cpu_cores"] for n in post.get("node_resources", []))
    total_base_mem = sum(n["mem_mib"]   for n in baseline.get("node_resources", []))
    total_post_mem = sum(n["mem_mib"]   for n in post.get("node_resources", []))
    kac_cpu        = sum(p["cpu_cores"] for p in post.get("kac_pods", []))
    kac_mem        = sum(p["mem_mib"]   for p in post.get("kac_pods", []))

    delta["summary"] = {
        "cluster_cpu_base":      round(total_base_cpu, 4),
        "cluster_cpu_post":      round(total_post_cpu, 4),
        "cluster_cpu_delta":     round(total_post_cpu - total_base_cpu, 4),
        "cluster_cpu_pct":       round((total_post_cpu - total_base_cpu) / max(total_base_cpu, 0.001) * 100, 1),
        "cluster_mem_base":      round(total_base_mem, 2),
        "cluster_mem_post":      round(total_post_mem, 2),
        "cluster_mem_delta":     round(total_post_mem - total_base_mem, 2),
        "cluster_mem_pct":       round((total_post_mem - total_base_mem) / max(total_base_mem, 0.001) * 100, 1),
        "kac_cpu_cores":         round(kac_cpu, 4),
        "kac_mem_mib":           round(kac_mem, 2),
        "kac_pod_count":         len(post.get("kac_pods", [])),
    }

    # Latency delta
    base_lat = baseline.get("admission_latency", {})
    post_lat  = post.get("admission_latency", {})
    delta["latency_delta"] = {
        "p50": _safe_diff(post_lat.get("p50"), base_lat.get("p50")),
        "p95": _safe_diff(post_lat.get("p95"), base_lat.get("p95")),
        "p99": _safe_diff(post_lat.get("p99"), base_lat.get("p99")),
    }
    return delta


# ─────────────────────────────────────────────────────────────────────────────
# RISK ASSESSMENT
# ─────────────────────────────────────────────────────────────────────────────
def assess_risk(delta, post):
    findings, risk_level = [], "LOW"
    s   = delta.get("summary", {})
    dep = post.get("kac_deployment", {})
    lat = post.get("admission_latency", {})

    # Deployment health
    desired  = dep.get("desired", 0)
    ready    = dep.get("ready", 0)
    unavail  = dep.get("unavailable", 0)
    if desired > 0 and ready < desired:
        risk_level = "HIGH"
        findings.append(f"CRIT  KAC Deployment: only {ready}/{desired} replicas ready ({unavail} unavailable).")
    elif desired > 0:
        findings.append(f"OK    KAC Deployment: {ready}/{desired} replicas ready.")
    if ready < 2:
        findings.append("WARN  Running fewer than 2 KAC replicas — no HA redundancy.")
        risk_level = "MEDIUM" if risk_level == "LOW" else risk_level

    # Cluster CPU overhead
    cpu_pct = s.get("cluster_cpu_pct", 0)
    if cpu_pct > 20:
        risk_level = "HIGH"
        findings.append(f"CRIT  Cluster CPU increased {cpu_pct}% after KAC install — check sensor limits.")
    elif cpu_pct > 10:
        findings.append(f"WARN  Cluster CPU increased {cpu_pct}% — monitor under peak load.")
        risk_level = "MEDIUM" if risk_level == "LOW" else risk_level
    else:
        findings.append(f"OK    Cluster CPU overhead {cpu_pct}% — acceptable.")

    # Cluster memory overhead
    mem_pct = s.get("cluster_mem_pct", 0)
    if mem_pct > 15:
        risk_level = "HIGH"
        findings.append(f"CRIT  Cluster memory increased {mem_pct}% — validate node headroom.")
    elif mem_pct > 8:
        findings.append(f"WARN  Cluster memory increased {mem_pct}% — review limits.")
        risk_level = "MEDIUM" if risk_level == "LOW" else risk_level
    else:
        findings.append(f"OK    Cluster memory overhead {mem_pct}% — acceptable.")

    # Webhook config
    for wh in post.get("webhook_config", {}).get("webhooks", []):
        fp      = wh.get("failurePolicy", "Fail")
        timeout = wh.get("timeoutSeconds", 10)
        if fp == "Fail":
            findings.append(
                f"WARN  Webhook '{wh['name']}' failurePolicy=Fail / timeoutSeconds={timeout} — "
                "KAC outage blocks pod scheduling cluster-wide."
            )
            risk_level = "MEDIUM" if risk_level == "LOW" else risk_level
        else:
            findings.append(f"OK    Webhook '{wh['name']}' failurePolicy=Ignore — safe failure mode.")

    # Latency
    p99 = lat.get("p99")
    if p99 and p99 > 1000:
        risk_level = "HIGH"
        findings.append(f"CRIT  KAC p99 latency {p99}ms — exceeds 1000ms critical threshold.")
    elif p99 and p99 > 500:
        findings.append(f"WARN  KAC p99 latency {p99}ms — approaching 1000ms threshold.")
        risk_level = "MEDIUM" if risk_level == "LOW" else risk_level
    elif p99:
        findings.append(f"OK    KAC p99 latency {p99}ms — healthy.")

    # Restart counts
    for pod in post.get("kac_pods", []):
        if pod.get("restarts", 0) > 5:
            risk_level = "MEDIUM" if risk_level == "LOW" else risk_level
            findings.append(f"WARN  Pod {pod['pod']} has {pod['restarts']} restarts — check OOMKill or probe failures.")

    return {"risk_level": risk_level, "findings": findings}


# ─────────────────────────────────────────────────────────────────────────────
# CHARTS
# ─────────────────────────────────────────────────────────────────────────────
def generate_charts(post, delta, output_dir):
    if not HAS_PLOTLY:
        print("  [WARN] plotly not installed — skipping charts.")
        return []
    os.makedirs(output_dir, exist_ok=True)
    charts = []

    # 1. Node CPU before vs after
    nodes = delta.get("nodes", [])
    if nodes:
        names = [n["name"] for n in nodes]
        fig = go.Figure()
        fig.add_trace(go.Bar(name="Baseline",    x=names, y=[n["cpu_base"] for n in nodes]))
        fig.add_trace(go.Bar(name="Post-Deploy", x=names, y=[n["cpu_post"] for n in nodes]))
        fig.update_layout(
            title={"text": "Node CPU: Baseline vs Post-Deploy<br>"
                           "<span style='font-size:14px;font-weight:normal;'>"
                           "CPU cores used per node</span>"},
            barmode="group",
            legend=dict(orientation="h", yanchor="bottom", y=1.05, xanchor="center", x=0.5),
        )
        fig.update_xaxes(title_text="Node")
        fig.update_yaxes(title_text="CPU (cores)")
        path = os.path.join(output_dir, "chart_node_cpu.png")
        fig.write_image(path)
        charts.append(path)

    # 2. Node Memory before vs after
    if nodes:
        fig = go.Figure()
        fig.add_trace(go.Bar(name="Baseline",    x=names, y=[n["mem_base"] for n in nodes]))
        fig.add_trace(go.Bar(name="Post-Deploy", x=names, y=[n["mem_post"] for n in nodes]))
        fig.update_layout(
            title={"text": "Node Memory: Baseline vs Post-Deploy<br>"
                           "<span style='font-size:14px;font-weight:normal;'>"
                           "MiB used per node</span>"},
            barmode="group",
            legend=dict(orientation="h", yanchor="bottom", y=1.05, xanchor="center", x=0.5),
        )
        fig.update_xaxes(title_text="Node")
        fig.update_yaxes(title_text="Memory (MiB)")
        path = os.path.join(output_dir, "chart_node_mem.png")
        fig.write_image(path)
        charts.append(path)

    # 3. KAC pod resource breakdown
    kac_pods = post.get("kac_pods", [])
    if kac_pods:
        pod_names = [p["pod"][:28] for p in kac_pods]
        fig = make_subplots(rows=1, cols=2, subplot_titles=("CPU (cores)", "Memory (MiB)"))
        fig.add_trace(go.Bar(x=pod_names, y=[p["cpu_cores"] for p in kac_pods], name="CPU"),    row=1, col=1)
        fig.add_trace(go.Bar(x=pod_names, y=[p["mem_mib"]   for p in kac_pods], name="Memory"), row=1, col=2)
        fig.update_layout(
            title={"text": "KAC Pod Resource Usage (Post-Deploy)<br>"
                           "<span style='font-size:14px;font-weight:normal;'>"
                           "Per-pod CPU and memory</span>"},
            showlegend=False,
        )
        fig.update_xaxes(tickangle=-20)
        path = os.path.join(output_dir, "chart_kac_pods.png")
        fig.write_image(path)
        charts.append(path)

    # 4. Per-container breakdown (falcon-ac vs falcon-client)
    container_rows = []
    for pod in kac_pods:
        for c in pod.get("containers", []):
            container_rows.append({
                "label": f"{pod['pod'][:18]}/{c['container']}",
                "cpu":   c["cpu_cores"],
                "mem":   c["mem_mib"],
            })
    if container_rows:
        labels = [r["label"] for r in container_rows]
        fig = go.Figure()
        fig.add_trace(go.Bar(name="CPU (cores)", x=labels, y=[r["cpu"] for r in container_rows]))
        fig.add_trace(go.Bar(name="Mem (MiB)",   x=labels, y=[r["mem"] for r in container_rows]))
        fig.update_layout(
            title={"text": "KAC Container Breakdown (falcon-ac vs falcon-client)<br>"
                           "<span style='font-size:14px;font-weight:normal;'>"
                           "Per-container CPU and memory</span>"},
            barmode="group",
            legend=dict(orientation="h", yanchor="bottom", y=1.05, xanchor="center", x=0.5),
        )
        fig.update_xaxes(title_text="Pod / Container", tickangle=-20)
        fig.update_yaxes(title_text="Usage")
        path = os.path.join(output_dir, "chart_kac_containers.png")
        fig.write_image(path)
        charts.append(path)

    # 5. Webhook latency
    lat = post.get("admission_latency", {})
    lat_vals = {k: v for k, v in lat.items() if k in ("p50", "p95", "p99") and v is not None}
    if lat_vals:
        fig = go.Figure(go.Bar(
            x=list(lat_vals.keys()),
            y=list(lat_vals.values()),
            text=[f"{v} ms" for v in lat_vals.values()],
            textposition="outside",
        ))
        fig.add_hline(y=500,  line_dash="dash", line_color="#ff9900",
                      annotation_text="500ms warn",  annotation_position="top right")
        fig.add_hline(y=1000, line_dash="dot",  line_color="#cc0000",
                      annotation_text="1000ms crit", annotation_position="top right")
        fig.update_layout(
            title={"text": "KAC Webhook Latency<br>"
                           "<span style='font-size:14px;font-weight:normal;'>"
                           "p50 / p95 / p99 ms from apiserver /metrics</span>"},
        )
        fig.update_xaxes(title_text="Percentile")
        fig.update_yaxes(title_text="Latency (ms)")
        path = os.path.join(output_dir, "chart_kac_latency.png")
        fig.write_image(path)
        charts.append(path)

    return charts


# ─────────────────────────────────────────────────────────────────────────────
# DOCX REPORT
# ─────────────────────────────────────────────────────────────────────────────
def _cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _add_table(doc, headers, rows, hdr_fill="1F3864", hdr_font="FFFFFF"):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(
            int(hdr_font[0:2],16), int(hdr_font[2:4],16), int(hdr_font[4:6],16))
        _cell_bg(hdr_cells[i], hdr_fill)
    for ridx, row in enumerate(rows):
        rc = tbl.add_row().cells
        fill = "EEF0F5" if ridx % 2 == 0 else "FFFFFF"
        for cidx, val in enumerate(row):
            rc[cidx].text = str(val)
            _cell_bg(rc[cidx], fill)
    doc.add_paragraph()


def generate_docx_report(baseline, post, delta, assessment, charts, output_path):
    if not HAS_DOCX:
        print("[ERROR] python-docx not installed — skipping DOCX."); return

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    ci  = post.get("cluster_info",  {})
    hr  = post.get("helm_release",  {})
    dep = post.get("kac_deployment",{})
    lat = post.get("admission_latency", {})
    wh  = post.get("webhook_config",{})
    s   = delta.get("summary",      {})

    # Title
    t = doc.add_heading("CrowdStrike Falcon KAC", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("Performance Impact Assessment Report")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(14); p.runs[0].bold = True
    doc.add_paragraph()

    # 1. Cluster & Helm Identity
    doc.add_heading("1. Cluster & Helm Release Identity", level=1)
    _add_table(doc, ["Field", "Value"], [
        ["Cluster Name",          ci.get("cluster_name",  "unknown")],
        ["Kubeconfig Context",    ci.get("context",        "unknown")],
        ["API Server URL",        ci.get("server_url",     "unknown")],
        ["Platform",              ci.get("k8s_platform",   "unknown")],
        ["Kubernetes Version",    ci.get("k8s_version",    "unknown")],
        ["Total / Worker Nodes",  f"{ci.get('node_count','N/A')} / {ci.get('worker_count','N/A')}"],
        ["Helm Release",          hr.get("release",        "unknown")],
        ["Helm Namespace",        hr.get("namespace",      "unknown")],
        ["Chart",                 f"{hr.get('chart','?')}-{hr.get('chart_version','?')}"],
        ["App Version",           hr.get("app_version",    "unknown")],
        ["Helm Status",           hr.get("status",         "unknown")],
        ["Last Deployed",         hr.get("deployed_at",    "unknown")],
        ["Baseline Captured",     baseline.get("timestamp","N/A")],
        ["Post-Deploy Captured",  post.get("timestamp",    "N/A")],
        ["Report Generated",      datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")],
    ])

    # 2. Risk Assessment
    doc.add_heading("2. Risk Assessment", level=1)
    risk = assessment["risk_level"]
    risk_rgb = {"LOW": (26,127,55), "MEDIUM": (154,103,0), "HIGH": (185,28,28)}
    p = doc.add_paragraph()
    run = p.add_run(f"Overall Risk Level:  {risk}")
    run.bold = True; run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(*risk_rgb.get(risk, (0,0,0)))
    for f in assessment["findings"]:
        doc.add_paragraph(f, style="List Bullet")
    doc.add_paragraph()

    # 3. Resource Delta
    doc.add_heading("3. Resource Delta Summary", level=1)
    _add_table(doc, ["Metric", "Baseline", "Post-Deploy", "Delta", "% Change"], [
        ["Cluster CPU (cores)",
         f"{s.get('cluster_cpu_base',0):.3f}", f"{s.get('cluster_cpu_post',0):.3f}",
         f"{s.get('cluster_cpu_delta',0):+.3f}", f"{s.get('cluster_cpu_pct',0):+.1f}%"],
        ["Cluster Memory (MiB)",
         f"{s.get('cluster_mem_base',0):.0f}", f"{s.get('cluster_mem_post',0):.0f}",
         f"{s.get('cluster_mem_delta',0):+.0f}", f"{s.get('cluster_mem_pct',0):+.1f}%"],
        ["KAC CPU (cores)", "N/A", f"{s.get('kac_cpu_cores',0):.3f}", "–", "–"],
        ["KAC Memory (MiB)", "N/A", f"{s.get('kac_mem_mib',0):.0f}", "–", "–"],
        ["KAC Pod Count", "N/A", str(s.get("kac_pod_count",0)), "–", "–"],
    ])

    # 4. Per-node detail
    doc.add_heading("4. Per-Node Resource Detail", level=1)
    _add_table(doc,
        ["Node", "CPU Base", "CPU Post", "CPU Δ", "Mem Base (MiB)", "Mem Post (MiB)", "Mem Δ (MiB)"],
        [[n["name"],
          f"{n['cpu_base']:.3f}", f"{n['cpu_post']:.3f}", f"{n['cpu_delta']:+.3f}",
          f"{n['mem_base']:.0f}", f"{n['mem_post']:.0f}", f"{n['mem_delta_mib']:+.0f}"]
         for n in delta.get("nodes", [])]
    )

    # 5. KAC pod + container detail
    doc.add_heading("5. KAC Pod & Container Detail (Post-Deploy)", level=1)
    container_rows = []
    for pod in post.get("kac_pods", []):
        for c in pod.get("containers", []):
            container_rows.append([
                pod["pod"], c["container"],
                f"{c['cpu_cores']:.3f}", f"{c['mem_mib']:.0f}",
                str(pod.get("restarts", 0)), pod.get("node",""),
            ])
    if container_rows:
        _add_table(doc,
            ["Pod", "Container", "CPU (cores)", "Memory (MiB)", "Restarts", "Node"],
            container_rows)
    else:
        doc.add_paragraph("No KAC pods found (metrics-server may not be installed).")

    # 6. Deployment status
    doc.add_heading("6. KAC Deployment Status", level=1)
    cond_str = ", ".join(f"{c['type']}={c['status']}" for c in dep.get("conditions", []))
    _add_table(doc, ["Field", "Value"], [
        ["Namespace",   dep.get("namespace",  "")],
        ["Desired",     str(dep.get("desired",  ""))],
        ["Ready",       str(dep.get("ready",    ""))],
        ["Available",   str(dep.get("available",""))],
        ["Unavailable", str(dep.get("unavailable",""))],
        ["Conditions",  cond_str],
    ])
    doc.add_heading("6.1  Container Resource Requests / Limits", level=2)
    cr_rows = []
    for cname, res in dep.get("container_resources", {}).items():
        reqs, lims = res.get("requests",{}), res.get("limits",{})
        cr_rows.append([cname,
            reqs.get("cpu","–"), reqs.get("memory","–"),
            lims.get("cpu","–"), lims.get("memory","–")])
    if cr_rows:
        _add_table(doc, ["Container","Req CPU","Req Memory","Limit CPU","Limit Memory"], cr_rows)

    # 7. Helm values
    vs = hr.get("values_summary", {})
    if vs:
        doc.add_heading("7. Helm Values (Performance-Relevant)", level=1)
        _add_table(doc, ["Parameter", "Value"], [
            ["replicas",               str(vs.get("replicas",""))],
            ["webhook.failurePolicy",  str(vs.get("failurePolicy",""))],
            ["webhook.timeoutSeconds", str(vs.get("timeoutSeconds",""))],
            ["priorityClassName",      str(vs.get("priorityClassName",""))],
            ["resources.falcon-ac",    str(vs.get("resources_ac",""))],
            ["resources.falcon-client",str(vs.get("resources_client",""))],
        ])

    # 8. Webhook config
    doc.add_heading("8. ValidatingWebhookConfiguration", level=1)
    if wh:
        doc.add_paragraph(f"Name: {wh.get('name','')}")
        _add_table(doc,
            ["Webhook","Failure Policy","Timeout (s)","Match Policy","SideEffects","NS Selector","Rules"],
            [[h["name"], h["failurePolicy"], str(h["timeoutSeconds"]),
              h["matchPolicy"], h["sideEffects"],
              "Yes" if h["namespaceSelector"] else "No", str(h["rules"])]
             for h in wh.get("webhooks", [])])
    else:
        doc.add_paragraph("No Falcon ValidatingWebhookConfiguration found.")

    # 9. Admission latency
    doc.add_heading("9. Admission Webhook Latency", level=1)
    lat_rows = []
    for k, lbl in [("p50","p50 (ms)"),("p95","p95 (ms)"),("p99","p99 (ms)"),("mean_ms","Mean (ms)")]:
        v = lat.get(k)
        if v is not None:
            lat_rows.append([lbl, str(v)])
    lat_rows.append(["Sample Count", str(lat.get("sample_count",0))])
    _add_table(doc, ["Metric","Value"], lat_rows)

    # 10. HPA
    hpa = post.get("hpa", {})
    if hpa:
        doc.add_heading("10. Horizontal Pod Autoscaler", level=1)
        _add_table(doc, ["Field","Value"], [
            ["Min Replicas",  str(hpa.get("min_replicas",""))],
            ["Max Replicas",  str(hpa.get("max_replicas",""))],
            ["Current",       str(hpa.get("current",""))],
            ["Desired",       str(hpa.get("desired",""))],
            ["Metrics",       ", ".join(hpa.get("metrics",[]))],
        ])

    # 11. ResourceQuota
    rq = post.get("resource_quota", {})
    if rq:
        doc.add_heading("11. ResourceQuota", level=1)
        rq_rows = []
        for qname, qdata in rq.items():
            for resource, hard_val in qdata.get("hard", {}).items():
                used_val = qdata.get("used", {}).get(resource, "–")
                rq_rows.append([qname, resource, hard_val, used_val])
        if rq_rows:
            _add_table(doc, ["Quota Name","Resource","Hard Limit","Used"], rq_rows)

    # 12. Events
    events = post.get("events", [])
    if events:
        doc.add_heading("12. Events (falcon-kac namespace)", level=1)
        _add_table(doc,
            ["Type","Reason","Count","Last Seen","Object","Message"],
            [[e["type"], e["reason"], str(e["count"]),
              e["last_ts"], e["object"], e["message"]] for e in events])

    # 13. Charts
    if charts:
        doc.add_heading("13. Performance Charts", level=1)
        titles = {
            "chart_node_cpu.png":        "Figure 1 – Node CPU: Baseline vs Post-Deploy",
            "chart_node_mem.png":        "Figure 2 – Node Memory: Baseline vs Post-Deploy",
            "chart_kac_pods.png":        "Figure 3 – KAC Pod Resource Usage",
            "chart_kac_containers.png":  "Figure 4 – KAC Container Breakdown (falcon-ac vs falcon-client)",
            "chart_kac_latency.png":     "Figure 5 – KAC Webhook Latency p50/p95/p99",
        }
        for cp in charts:
            if os.path.exists(cp):
                doc.add_paragraph(titles.get(os.path.basename(cp), os.path.basename(cp))).runs[0].bold = True
                doc.add_picture(cp, width=Inches(6.0))
                doc.add_paragraph()

    # 14. Recommendations
    doc.add_heading("14. Recommended Next Steps", level=1)
    for rec in [
        "Scale to >=2 replicas for HA:  helm upgrade falcon-kac ... --set replicas=2",
        "If p99 > 500ms, increase falcon-ac CPU limits in values.yaml under resources.",
        "Set webhook.failurePolicy=Ignore and webhook.timeoutSeconds=10 for availability.",
        "Add priorityClassName: system-cluster-critical to prevent eviction.",
        "Enable HPA targeting 60% CPU to auto-scale under admission burst.",
        "Monitor container_cpu_cfs_throttled_periods_total{container='falcon-ac'} for throttling.",
        "Re-run during peak load for representative admission latency measurements.",
    ]:
        doc.add_paragraph(rec, style="List Number")

    # Footer
    doc.add_paragraph()
    fp = doc.add_paragraph(
        f"falcon_kac_perf_assessment.py  |  "
        f"Cluster: {ci.get('cluster_name','unknown')}  |  "
        f"Release: {hr.get('release','unknown')} ({hr.get('chart_version','?')})  |  "
        f"{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )
    fp.runs[0].font.size = Pt(9)
    fp.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(output_path)
    print(f"  DOCX  -> {output_path}")


# ─────────────────────────────────────────────────────────────────────────────
# HTML REPORT
# ─────────────────────────────────────────────────────────────────────────────
HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>Falcon KAC – Performance Impact Report</title>
<style>
  body{font-family:'Segoe UI',Arial,sans-serif;background:#0f1117;color:#e0e0e0;margin:40px}
  h1{color:#e85c27;font-size:2em} h2{color:#ff7b4f;border-bottom:1px solid #333;padding-bottom:6px;margin-top:28px}
  h3{color:#c0c0c0;margin-top:16px}
  table{width:100%;border-collapse:collapse;margin-bottom:20px}
  th{background:#1e2230;color:#ff7b4f;padding:8px 12px;text-align:left}
  td{padding:7px 12px;border-bottom:1px solid #2a2d3a}
  tr:nth-child(even) td{background:#181c29}
  .badge-low{background:#1a7f37;color:#fff;padding:2px 12px;border-radius:12px;font-weight:bold}
  .badge-medium{background:#9a6700;color:#fff;padding:2px 12px;border-radius:12px;font-weight:bold}
  .badge-high{background:#b91c1c;color:#fff;padding:2px 12px;border-radius:12px;font-weight:bold}
  .finding{padding:6px 0;border-bottom:1px dashed #2a2d3a;font-family:monospace}
  .ok{color:#4caf50}.warn{color:#ff9800}.crit{color:#f44336}
  .chart{max-width:100%;margin:16px 0;border:1px solid #333;border-radius:8px}
  .section{background:#13161f;border-radius:8px;padding:16px 24px;margin-bottom:24px}
  .meta{color:#888;font-size:0.85em} code{background:#1a1d2b;padding:2px 6px;border-radius:4px;font-family:monospace}
</style>
</head>
<body>
<h1>&#x1F6E1;&#xFE0F; CrowdStrike Falcon KAC &#8211; Performance Impact Report</h1>
<p class="meta">Generated: {{ timestamp }}</p>

<div class="section">
  <h2>Cluster &amp; Helm Release Identity</h2>
  <table>
    <tr><th>Field</th><th>Value</th></tr>
    <tr><td>Cluster Name</td><td><strong>{{ ci.cluster_name }}</strong></td></tr>
    <tr><td>Kubeconfig Context</td><td>{{ ci.context }}</td></tr>
    <tr><td>API Server URL</td><td><code>{{ ci.server_url }}</code></td></tr>
    <tr><td>Platform</td><td>{{ ci.k8s_platform }}</td></tr>
    <tr><td>Kubernetes Version</td><td>{{ ci.k8s_version }}</td></tr>
    <tr><td>Total / Worker Nodes</td><td>{{ ci.node_count | default('N/A') }} / {{ ci.worker_count | default('N/A') }}</td></tr>
    <tr><td>Helm Release</td><td><strong>{{ hr.release }}</strong></td></tr>
    <tr><td>Chart</td><td>{{ hr.chart }}-{{ hr.chart_version }}</td></tr>
    <tr><td>App Version</td><td>{{ hr.app_version }}</td></tr>
    <tr><td>Helm Status</td><td>{{ hr.status }}</td></tr>
    <tr><td>Last Deployed</td><td>{{ hr.deployed_at }}</td></tr>
    <tr><td>Baseline Captured</td><td>{{ baseline_ts }}</td></tr>
    <tr><td>Post-Deploy Captured</td><td>{{ post_ts }}</td></tr>
  </table>
</div>

<div class="section">
  <h2>Risk Assessment</h2>
  <p>Risk Level: <span class="badge-{{ risk_level | lower }}">{{ risk_level }}</span></p>
  {% for f in findings %}
  <div class="finding {% if f.startswith('CRIT') %}crit{% elif f.startswith('WARN') %}warn{% else %}ok{% endif %}">{{ f }}</div>
  {% endfor %}
</div>

<div class="section">
  <h2>Resource Delta Summary</h2>
  <table>
    <tr><th>Metric</th><th>Baseline</th><th>Post-Deploy</th><th>Delta</th><th>% Change</th></tr>
    <tr><td>Cluster CPU (cores)</td>
        <td>{{ "%.3f"|format(s.cluster_cpu_base) }}</td><td>{{ "%.3f"|format(s.cluster_cpu_post) }}</td>
        <td>{{ "%+.3f"|format(s.cluster_cpu_delta) }}</td><td>{{ "%+.1f"|format(s.cluster_cpu_pct) }}%</td></tr>
    <tr><td>Cluster Memory (MiB)</td>
        <td>{{ "%.0f"|format(s.cluster_mem_base) }}</td><td>{{ "%.0f"|format(s.cluster_mem_post) }}</td>
        <td>{{ "%+.0f"|format(s.cluster_mem_delta) }}</td><td>{{ "%+.1f"|format(s.cluster_mem_pct) }}%</td></tr>
    <tr><td>KAC CPU (cores)</td><td>N/A</td><td>{{ "%.3f"|format(s.kac_cpu_cores) }}</td><td>&#8211;</td><td>&#8211;</td></tr>
    <tr><td>KAC Memory (MiB)</td><td>N/A</td><td>{{ "%.0f"|format(s.kac_mem_mib) }}</td><td>&#8211;</td><td>&#8211;</td></tr>
    <tr><td>KAC Pod Count</td><td>N/A</td><td>{{ s.kac_pod_count }}</td><td>&#8211;</td><td>&#8211;</td></tr>
  </table>
</div>

<div class="section">
  <h2>Per-Node Resource Detail</h2>
  <table>
    <tr><th>Node</th><th>CPU Base</th><th>CPU Post</th><th>CPU &#916;</th>
        <th>Mem Base (MiB)</th><th>Mem Post (MiB)</th><th>Mem &#916; (MiB)</th></tr>
    {% for n in node_delta %}
    <tr><td>{{ n.name }}</td>
        <td>{{ "%.3f"|format(n.cpu_base) }}</td><td>{{ "%.3f"|format(n.cpu_post) }}</td>
        <td>{{ "%+.3f"|format(n.cpu_delta) }}</td>
        <td>{{ "%.0f"|format(n.mem_base) }}</td><td>{{ "%.0f"|format(n.mem_post) }}</td>
        <td>{{ "%+.0f"|format(n.mem_delta_mib) }}</td></tr>
    {% endfor %}
  </table>
</div>

<div class="section">
  <h2>KAC Pod &amp; Container Detail</h2>
  <table>
    <tr><th>Pod</th><th>Container</th><th>CPU (cores)</th><th>Memory (MiB)</th><th>Restarts</th><th>Node</th></tr>
    {% for pod in kac_pods %}{% for c in pod.containers %}
    <tr><td>{{ pod.pod }}</td><td>{{ c.container }}</td>
        <td>{{ "%.3f"|format(c.cpu_cores) }}</td><td>{{ "%.0f"|format(c.mem_mib) }}</td>
        <td>{{ pod.restarts }}</td><td>{{ pod.node }}</td></tr>
    {% endfor %}{% endfor %}
  </table>
</div>

<div class="section">
  <h2>KAC Deployment Status</h2>
  {% if dep %}
  <table>
    <tr><th>Namespace</th><th>Desired</th><th>Ready</th><th>Available</th><th>Unavailable</th></tr>
    <tr><td>{{ dep.namespace }}</td><td>{{ dep.desired }}</td><td>{{ dep.ready }}</td>
        <td>{{ dep.available }}</td><td>{{ dep.unavailable }}</td></tr>
  </table>
  <h3>Container Resource Requests / Limits</h3>
  <table>
    <tr><th>Container</th><th>Req CPU</th><th>Req Memory</th><th>Limit CPU</th><th>Limit Memory</th></tr>
    {% for cname, res in dep.container_resources.items() %}
    <tr><td>{{ cname }}</td>
        <td>{{ res.requests.get('cpu','–') }}</td><td>{{ res.requests.get('memory','–') }}</td>
        <td>{{ res.limits.get('cpu','–') }}</td><td>{{ res.limits.get('memory','–') }}</td></tr>
    {% endfor %}
  </table>
  {% endif %}
</div>

{% if hr.values_summary %}
<div class="section">
  <h2>Helm Values (Performance-Relevant)</h2>
  <table>
    <tr><th>Parameter</th><th>Value</th></tr>
    <tr><td>replicas</td><td>{{ hr.values_summary.replicas }}</td></tr>
    <tr><td>webhook.failurePolicy</td><td>{{ hr.values_summary.failurePolicy }}</td></tr>
    <tr><td>webhook.timeoutSeconds</td><td>{{ hr.values_summary.timeoutSeconds }}</td></tr>
    <tr><td>priorityClassName</td><td>{{ hr.values_summary.priorityClassName }}</td></tr>
    <tr><td>resources.falcon-ac</td><td><code>{{ hr.values_summary.resources_ac }}</code></td></tr>
    <tr><td>resources.falcon-client</td><td><code>{{ hr.values_summary.resources_client }}</code></td></tr>
  </table>
</div>
{% endif %}


{% if hr.source == 'FluxCD HelmRelease' %}
<div class="section">
  <h2>&#x1F4E6; FluxCD Deployment Info</h2>
  <table>
    <tr><th>Field</th><th>Value</th></tr>
    <tr><td>Source</td><td>FluxCD HelmRelease</td></tr>
    <tr><td>Git Source</td><td><code>{{ hr.git_source }}</code></td></tr>
    <tr><td>Chart Path</td><td><code>{{ hr.chart_path }}</code></td></tr>
    <tr><td>Version Selector</td><td><code>{{ hr.chart_selector }}</code></td></tr>
    <tr><td>Flux Ready</td><td>{{ "✅ True" if hr.flux_ready else "❌ False" }}</td></tr>
    <tr><td>Flux Reason</td><td>{{ hr.flux_reason }}</td></tr>
    <tr><td>Flux Message</td><td>{{ hr.flux_message }}</td></tr>
    <tr><td>Last Reconcile</td><td>{{ hr.reconcile_at }}</td></tr>
  </table>
  {% if flux_helmchart %}
  <h3>HelmChart CR (flux-system/{{ flux_helmchart.name }})</h3>
  <table>
    <tr><th>Field</th><th>Value</th></tr>
    <tr><td>Chart</td><td>{{ flux_helmchart.chart }}</td></tr>
    <tr><td>Version</td><td>{{ flux_helmchart.version }}</td></tr>
    <tr><td>Source Kind</td><td>{{ flux_helmchart.source_kind }}</td></tr>
    <tr><td>Source Name</td><td>{{ flux_helmchart.source_name }}</td></tr>
    <tr><td>Observed Version</td><td>{{ flux_helmchart.observed_version }}</td></tr>
    <tr><td>Artifact Path</td><td><code>{{ flux_helmchart.artifact_path }}</code></td></tr>
    <tr><td>Revision</td><td><code>{{ flux_helmchart.revision }}</code></td></tr>
    <tr><td>Ready</td><td>{{ "✅ True" if flux_helmchart.ready else "❌ False" }}</td></tr>
  </table>
  {% endif %}
</div>
{% endif %}
<div class="section">
  <h2>Webhook Configuration</h2>
  {% if webhook %}
  <p>Name: <code>{{ webhook.name }}</code></p>
  <table>
    <tr><th>Webhook</th><th>Failure Policy</th><th>Timeout (s)</th><th>Match Policy</th>
        <th>SideEffects</th><th>NS Selector</th><th>Rules</th></tr>
    {% for wh in webhook.webhooks %}
    <tr><td>{{ wh.name }}</td><td>{{ wh.failurePolicy }}</td><td>{{ wh.timeoutSeconds }}</td>
        <td>{{ wh.matchPolicy }}</td><td>{{ wh.sideEffects }}</td>
        <td>{{ "Yes" if wh.namespaceSelector else "No" }}</td><td>{{ wh.rules }}</td></tr>
    {% endfor %}
  </table>
  {% else %}<p class="meta">No Falcon webhook found.</p>{% endif %}
</div>

<div class="section">
  <h2>Admission Webhook Latency</h2>
  <table><tr><th>Metric</th><th>Value</th></tr>
    {% if lat.p50   is not none %}<tr><td>p50 (ms)</td><td>{{ lat.p50 }}</td></tr>{% endif %}
    {% if lat.p95   is not none %}<tr><td>p95 (ms)</td><td>{{ lat.p95 }}</td></tr>{% endif %}
    {% if lat.p99   is not none %}<tr><td>p99 (ms)</td><td>{{ lat.p99 }}</td></tr>{% endif %}
    {% if lat.mean_ms is not none %}<tr><td>Mean (ms)</td><td>{{ lat.mean_ms }}</td></tr>{% endif %}
    <tr><td>Sample Count</td><td>{{ lat.sample_count }}</td></tr>
  </table>
</div>

{% if hpa %}
<div class="section">
  <h2>Horizontal Pod Autoscaler</h2>
  <table><tr><th>Min</th><th>Max</th><th>Current</th><th>Desired</th><th>Metrics</th></tr>
    <tr><td>{{ hpa.min_replicas }}</td><td>{{ hpa.max_replicas }}</td>
        <td>{{ hpa.current }}</td><td>{{ hpa.desired }}</td>
        <td>{{ hpa.metrics | join(', ') }}</td></tr>
  </table>
</div>
{% endif %}

{% if events %}
<div class="section">
  <h2>Events (falcon-kac namespace)</h2>
  <table><tr><th>Type</th><th>Reason</th><th>Count</th><th>Last Seen</th><th>Object</th><th>Message</th></tr>
    {% for e in events %}
    <tr><td>{{ e.type }}</td><td>{{ e.reason }}</td><td>{{ e.count }}</td>
        <td>{{ e.last_ts }}</td><td>{{ e.object }}</td><td>{{ e.message }}</td></tr>
    {% endfor %}
  </table>
</div>
{% endif %}

{% if charts %}
<div class="section">
  <h2>Performance Charts</h2>
  {% for c in charts %}<img src="{{ c }}" class="chart"><br>{% endfor %}
</div>
{% endif %}

<div class="section">
  <h2>Recommended Next Steps</h2>
  <ol>
    <li>Scale to &#8805;2 replicas for HA: <code>helm upgrade falcon-kac ... --set replicas=2</code></li>
    <li>If p99 &gt; 500ms, increase <code>falcon-ac</code> CPU limits in values.yaml.</li>
    <li>Set <code>webhook.failurePolicy=Ignore</code> and <code>timeoutSeconds=10</code> for availability.</li>
    <li>Add <code>priorityClassName: system-cluster-critical</code> to prevent eviction.</li>
    <li>Enable HPA targeting 60% CPU to auto-scale under admission burst.</li>
    <li>Monitor <code>container_cpu_cfs_throttled_periods_total{container="falcon-ac"}</code> for throttling.</li>
    <li>Re-run during peak load for representative admission latency measurements.</li>
  </ol>
</div>
<p class="meta" style="text-align:center">
  falcon_kac_perf_assessment.py &nbsp;|&nbsp;
  Cluster: {{ ci.cluster_name }} &nbsp;|&nbsp;
  Release: {{ hr.release }} ({{ hr.chart_version }}) &nbsp;|&nbsp; {{ timestamp }}
</p>
</body></html>"""


def generate_html_report(baseline, post, delta, assessment, charts, output_path):
    if not HAS_JINJA:
        print("[ERROR] jinja2 not installed — skipping HTML."); return
    report_dir = os.path.dirname(os.path.abspath(output_path))
    rel_charts  = [os.path.relpath(c, report_dir) for c in charts]
    html = Template(HTML_TEMPLATE).render(
        timestamp   = datetime.datetime.now().isoformat(),
        baseline_ts = baseline.get("timestamp",""),
        post_ts     = post.get("timestamp",    ""),
        ci          = post.get("cluster_info", {}),
        hr          = post.get("helm_release", {}),
        risk_level  = assessment["risk_level"],
        findings    = assessment["findings"],
        s           = delta.get("summary",      {}),
        node_delta  = delta.get("nodes",        []),
        kac_pods    = post.get("kac_pods",      []),
        dep         = post.get("kac_deployment",{}),
        lat         = post.get("admission_latency",{}),
        webhook     = post.get("webhook_config",{}),
        hpa         = post.get("hpa",           {}),
        events      = post.get("events",        []),
        flux_helmchart = post.get("flux_helmchart", {}),
        charts      = rel_charts,
    )
    Path(output_path).write_text(html, encoding="utf-8")
    print(f"  HTML  -> {output_path}")


# ─────────────────────────────────────────────────────────────────────────────
# CLI
# ─────────────────────────────────────────────────────────────────────────────
def parse_args():
    p = argparse.ArgumentParser(
        description="Falcon KAC (Helm) Performance Impact Assessment",
        formatter_class=argparse.RawTextHelpFormatter,
        epilog="""
Workflow:
  Step 1 — Capture baseline BEFORE helm install (no falcon-kac required):
    python3 falcon_kac_perf_assessment.py --mode baseline --output output/baseline.json

  Step 2 — Deploy via FluxCD (HelmRelease + HelmChart in GitRepo):
    git commit + push your HelmRelease manifest targeting falcon-system namespace
    kubectl wait helmrelease/falcon-kac -n falcon-system --for=condition=Ready --timeout=5m

  Step 3 — Generate report:
    python3 falcon_kac_perf_assessment.py --mode post \\
        --baseline output/baseline.json \\
        --output output/kac_report \\
        --namespace falcon-system \\
        --release falcon-kac
        """,
    )
    p.add_argument("--mode",      required=True, choices=["baseline","post","sample"])
    p.add_argument("--baseline",  default=None)
    p.add_argument("--output",    required=True)
    p.add_argument("--namespace", default=DEFAULT_NAMESPACE,
                   help=f"KAC namespace (default: {DEFAULT_NAMESPACE})")
    p.add_argument("--release",   default=DEFAULT_RELEASE,
                   help=f"Helm/Flux release name (default: {DEFAULT_RELEASE})")
    p.add_argument("--samples",   type=int, default=6)
    p.add_argument("--interval",  type=int, default=60)
    p.add_argument("--chart-dir", default=None)
    p.add_argument("--wait",      type=int, default=120)
    return p.parse_args()


def main():
    args = parse_args()
    ns, rel = args.namespace, args.release

    if args.mode == "baseline":
        snap = collect_baseline_snapshot(ns, rel)
        Path(args.output).parent.mkdir(parents=True, exist_ok=True)
        Path(args.output).write_text(json.dumps(snap, indent=2))
        print(f"  Baseline -> {args.output}")

    elif args.mode == "sample":
        samples = collect_samples(args.samples, args.interval, ns, rel)
        Path(args.output).write_text(json.dumps(samples, indent=2))
        print(f"\n  Samples  -> {args.output}")

    elif args.mode == "post":
        if not args.baseline:
            print("[ERROR] --baseline is required for --mode post"); sys.exit(1)
        bpath = Path(args.baseline)
        if not bpath.exists():
            print(f"[ERROR] Baseline not found: {args.baseline}"); sys.exit(1)

        baseline  = json.loads(bpath.read_text())
        post      = collect_post_snapshot(ns, rel)
        delta     = compute_delta(baseline, post)
        assess    = assess_risk(delta, post)

        base_out  = str(args.output).removesuffix(".html").removesuffix(".docx")
        chart_dir = args.chart_dir or str(Path(base_out).parent)
        os.makedirs(chart_dir, exist_ok=True)
        charts    = generate_charts(post, delta, chart_dir)

        generate_html_report(baseline, post, delta, assess, charts, base_out + ".html")
        generate_docx_report(baseline, post, delta, assess, charts, base_out + ".docx")

        print("\n── Summary ─────────────────────────────────────────────────")
        ci = post.get("cluster_info",{})
        hr = post.get("helm_release",{})
        print(f"  Cluster  : {ci.get('cluster_name')} ({ci.get('k8s_platform')})")
        print(f"  Release  : {hr.get('release')} chart={hr.get('chart_version')} status={hr.get('status')}")
        print(f"  Risk     : {assess['risk_level']}")
        for f in assess["findings"]:
            print(f"  {f}")
        print("─────────────────────────────────────────────────────────────")


if __name__ == "__main__":
    main()
