#!/usr/bin/env python3
"""
falcon_kac_perf_assessment.py
Performance Impact Assessment – CrowdStrike Falcon KAC (Helm deployment)

Targets:
  - Deployment:  falcon-kac          (namespace: falcon-kac by default)
  - Pods:        falcon-kac-<hash>   (2 containers: falcon-ac + falcon-client)
  - Webhook:     ValidatingWebhookConfiguration with "falcon" in name
  - Helm labels: app.kubernetes.io/instance=<release>, helm.sh/chart=falcon-kac-*

Usage:
  # Step 1 – baseline BEFORE deploying
  python3 falcon_kac_perf_assessment.py --mode baseline --output /tmp/kac_baseline.json

  # Step 2 – post-deploy report (HTML + DOCX)
  python3 falcon_kac_perf_assessment.py --mode post \
      --baseline /tmp/kac_baseline.json \
      --output /tmp/kac_report \
      [--namespace falcon-kac] \
      [--release falcon-kac]

  # Continuous sampling
  python3 falcon_kac_perf_assessment.py --mode sample \
      --samples 10 --interval 30 --output /tmp/kac_samples.json

Dependencies:
  pip install jinja2 plotly kaleido python-docx
"""

import argparse
import json
import subprocess
import sys
import time
import datetime
import os
from pathlib import Path

try:
    import plotly.graph_objects as go
    HAS_PLOTLY = True
except ImportError:
    HAS_PLOTLY = False

try:
    from jinja2 import Template
    HAS_JINJA = True
except ImportError:
    HAS_JINJA = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ─────────────────────────────────────────────────────────────────────────────
# DEFAULTS  (override with CLI flags or --namespace / --release)
# ─────────────────────────────────────────────────────────────────────────────
DEFAULT_NAMESPACE    = "falcon-kac"
DEFAULT_RELEASE      = "falcon-kac"
DEPLOY_NAME          = "falcon-kac"
WEBHOOK_PATTERN      = "falcon"

# Container names inside the KAC pod (from CrowdStrike helm chart)
CONTAINER_AC_NAME    = "falcon-ac"       # admission controller container
CONTAINER_CLIENT     = "falcon-client"   # client/proxy container


# ─────────────────────────────────────────────────────────────────────────────
# KUBECTL HELPERS
# ─────────────────────────────────────────────────────────────────────────────
def run_kubectl(args: list, timeout: int = 30):
    cmd = ["kubectl"] + args
    try:
        r = subprocess.run(cmd, capture_output=True, text=True, timeout=timeout)
        return r.stdout, r.stderr, r.returncode
    except subprocess.TimeoutExpired:
        return "", f"kubectl timeout after {timeout}s", 1
    except FileNotFoundError:
        return "", "kubectl not found in PATH", 1


def kubectl_json(args: list):
    stdout, stderr, rc = run_kubectl(args + ["-o", "json"])
    if rc != 0:
        print(f"  [WARN] {stderr.strip()}", file=sys.stderr)
        return None
    try:
        return json.loads(stdout)
    except json.JSONDecodeError:
        return None


def kubectl_raw(path: str) -> str:
    stdout, _, _ = run_kubectl(["get", "--raw", path])
    return stdout


# ─────────────────────────────────────────────────────────────────────────────
# CLUSTER IDENTITY
# ─────────────────────────────────────────────────────────────────────────────
def collect_cluster_info() -> dict:
    info = {
        "cluster_name":  "unknown",
        "context":       "unknown",
        "server_url":    "unknown",
        "k8s_version":   "unknown",
        "k8s_platform":  "unknown",
        "node_count":    "N/A",
        "worker_count":  "N/A",
    }

    stdout, _, rc = run_kubectl(["config", "current-context"])
    if rc == 0:
        info["context"] = stdout.strip()

    cfg = kubectl_json(["config", "view"])
    if cfg:
        for ctx in cfg.get("contexts", []):
            if ctx.get("name") == info["context"]:
                info["cluster_name"] = ctx.get("context", {}).get("cluster", info["context"])
                break
        for cl in cfg.get("clusters", []):
            if cl.get("name") == info["cluster_name"]:
                info["server_url"] = cl.get("cluster", {}).get("server", "unknown")
                break

    ver = kubectl_json(["version"])
    if ver:
        sv  = ver.get("serverVersion", {})
        info["k8s_version"] = sv.get("gitVersion", "unknown")
        gv  = sv.get("gitVersion", "").lower()
        if   "eks"       in gv: info["k8s_platform"] = "Amazon EKS"
        elif "gke"       in gv: info["k8s_platform"] = "Google GKE"
        elif "aks"       in gv: info["k8s_platform"] = "Azure AKS"
        elif "openshift" in gv: info["k8s_platform"] = "OpenShift"
        else:
            nodes_obj = kubectl_json(["get", "nodes"])
            if nodes_obj:
                for n in nodes_obj.get("items", []):
                    pid = n.get("spec", {}).get("providerID", "")
                    if   "eks"   in pid: info["k8s_platform"] = "Amazon EKS";   break
                    elif "gce"   in pid: info["k8s_platform"] = "Google GKE";   break
                    elif "azure" in pid: info["k8s_platform"] = "Azure AKS";    break
                    elif n.get("metadata", {}).get("labels", {}).get("node.openshift.io/os_id"):
                        info["k8s_platform"] = "OpenShift"; break
                else:
                    info["k8s_platform"] = "Vanilla / On-Prem"

    nodes_obj = kubectl_json(["get", "nodes"])
    if nodes_obj:
        items = nodes_obj.get("items", [])
        info["node_count"]   = len(items)
        info["worker_count"] = sum(
            1 for n in items
            if not n.get("metadata", {}).get("labels", {}).get("node-role.kubernetes.io/control-plane")
            and not n.get("metadata", {}).get("labels", {}).get("node-role.kubernetes.io/master")
        )

    return info


# ─────────────────────────────────────────────────────────────────────────────
# HELM RELEASE INFO
# ─────────────────────────────────────────────────────────────────────────────
def collect_helm_release(namespace: str, release: str) -> dict:
    """
    Pull Helm release metadata via `helm status` and `helm get values`.
    Falls back gracefully if Helm CLI is not available.
    """
    info = {
        "release":     release,
        "namespace":   namespace,
        "chart":       "unknown",
        "chart_version": "unknown",
        "app_version": "unknown",
        "status":      "unknown",
        "deployed_at": "unknown",
        "values_summary": {},
    }

    # helm status
    try:
        r = subprocess.run(
            ["helm", "status", release, "-n", namespace, "--output", "json"],
            capture_output=True, text=True, timeout=20
        )
        if r.returncode == 0:
            hs = json.loads(r.stdout)
            info["chart"]         = hs.get("chart", {}).get("metadata", {}).get("name", "unknown")
            info["chart_version"] = hs.get("chart", {}).get("metadata", {}).get("version", "unknown")
            info["app_version"]   = hs.get("chart", {}).get("metadata", {}).get("appVersion", "unknown")
            info["status"]        = hs.get("info", {}).get("status", "unknown")
            info["deployed_at"]   = hs.get("info", {}).get("last_deployed", "unknown")
    except (FileNotFoundError, subprocess.TimeoutExpired, json.JSONDecodeError):
        pass

    # helm get values — extract resource config if set
    try:
        r = subprocess.run(
            ["helm", "get", "values", release, "-n", namespace, "--output", "json"],
            capture_output=True, text=True, timeout=20
        )
        if r.returncode == 0:
            vals = json.loads(r.stdout)
            # Surface the most performance-relevant values
            info["values_summary"] = {
                "replicas":          vals.get("replicas", "default"),
                "resources_ac":      vals.get("resources", {}).get("falcon-ac", {}),
                "resources_client":  vals.get("resources", {}).get("falcon-client", {}),
                "failurePolicy":     vals.get("webhook", {}).get("failurePolicy", "not set"),
                "timeoutSeconds":    vals.get("webhook", {}).get("timeoutSeconds", "not set"),
                "priorityClassName": vals.get("priorityClassName", "not set"),
            }
    except (FileNotFoundError, subprocess.TimeoutExpired, json.JSONDecodeError):
        pass

    return info


# ─────────────────────────────────────────────────────────────────────────────
# KAC-SPECIFIC METRIC COLLECTION
# ─────────────────────────────────────────────────────────────────────────────
def collect_kac_deployment(namespace: str) -> dict:
    """Full deployment object for falcon-kac."""
    dep = kubectl_json(["get", "deployment", DEPLOY_NAME, "-n", namespace])
    if not dep:
        return {}
    status = dep.get("status", {})
    spec   = dep.get("spec",   {})
    # Extract per-container resource requests/limits from pod spec
    containers = spec.get("template", {}).get("spec", {}).get("containers", [])
    container_resources = {}
    for c in containers:
        name = c.get("name", "unknown")
        res  = c.get("resources", {})
        container_resources[name] = {
            "requests": res.get("requests", {}),
            "limits":   res.get("limits",   {}),
        }
    return {
        "namespace":           namespace,
        "name":                DEPLOY_NAME,
        "desired":             spec.get("replicas", 1),
        "ready":               status.get("readyReplicas", 0),
        "available":           status.get("availableReplicas", 0),
        "updated":             status.get("updatedReplicas", 0),
        "unavailable":         status.get("unavailableReplicas", 0),
        "container_resources": container_resources,
        "conditions": [
            {"type": c["type"], "status": c["status"], "reason": c.get("reason", "")}
            for c in status.get("conditions", [])
        ],
    }


def collect_kac_pods(namespace: str) -> list:
    """
    Per-pod and per-container resource usage for falcon-kac pods.
    Uses kubectl top pods --containers for per-container breakdown.
    """
    pods = []

    # Pod-level totals
    stdout, _, rc = run_kubectl(["top", "pods", "-n", namespace, "--no-headers"])
    pod_totals = {}
    if rc == 0:
        for line in stdout.strip().splitlines():
            parts = line.split()
            if len(parts) >= 3:
                pod_totals[parts[0]] = {
                    "cpu_cores": _parse_cpu(parts[1]),
                    "mem_mib":   _parse_mem(parts[2]),
                }

    # Per-container breakdown
    stdout, _, rc = run_kubectl(
        ["top", "pods", "-n", namespace, "--containers", "--no-headers"]
    )
    pod_containers = {}
    if rc == 0:
        for line in stdout.strip().splitlines():
            parts = line.split()
            if len(parts) >= 4:
                pname, cname, cpu_raw, mem_raw = parts[0], parts[1], parts[2], parts[3]
                if pname not in pod_containers:
                    pod_containers[pname] = []
                pod_containers[pname].append({
                    "container": cname,
                    "cpu_cores": _parse_cpu(cpu_raw),
                    "mem_mib":   _parse_mem(mem_raw),
                })

    # Get pod metadata (node, restart counts, age)
    pods_obj = kubectl_json(["get", "pods", "-n", namespace,
                             "-l", f"app.kubernetes.io/instance={DEFAULT_RELEASE}"])
    pod_meta = {}
    if pods_obj:
        for p in pods_obj.get("items", []):
            pname = p["metadata"]["name"]
            cs    = p.get("status", {}).get("containerStatuses", [])
            pod_meta[pname] = {
                "node":     p.get("spec", {}).get("nodeName", "unknown"),
                "phase":    p.get("status", {}).get("phase", "unknown"),
                "restarts": sum(c.get("restartCount", 0) for c in cs),
                "age":      p.get("metadata", {}).get("creationTimestamp", ""),
            }

    # Merge
    all_pod_names = set(list(pod_totals.keys()) + list(pod_meta.keys()))
    for pname in all_pod_names:
        if not pname.startswith("falcon-kac"):
            continue
        totals = pod_totals.get(pname, {"cpu_cores": 0, "mem_mib": 0})
        meta   = pod_meta.get(pname, {})
        pods.append({
            "pod":        pname,
            "namespace":  namespace,
            "node":       meta.get("node", "unknown"),
            "phase":      meta.get("phase", "unknown"),
            "restarts":   meta.get("restarts", 0),
            "cpu_cores":  totals["cpu_cores"],
            "mem_mib":    totals["mem_mib"],
            "containers": pod_containers.get(pname, []),
        })

    return pods


def collect_webhook_config() -> dict:
    vwcs = kubectl_json(["get", "validatingwebhookconfigurations"])
    if not vwcs:
        return {}
    for item in vwcs.get("items", []):
        name = item.get("metadata", {}).get("name", "")
        if WEBHOOK_PATTERN in name.lower():
            webhooks = item.get("webhooks", [])
            return {
                "name": name,
                "webhooks": [
                    {
                        "name":             wh.get("name", ""),
                        "failurePolicy":    wh.get("failurePolicy", "unknown"),
                        "timeoutSeconds":   wh.get("timeoutSeconds", "default"),
                        "matchPolicy":      wh.get("matchPolicy", "Equivalent"),
                        "sideEffects":      wh.get("sideEffects", "None"),
                        "namespaceSelector": bool(wh.get("namespaceSelector")),
                        "objectSelector":   bool(wh.get("objectSelector")),
                        "rules":            len(wh.get("rules", [])),
                        "admissionReviewVersions": wh.get("admissionReviewVersions", []),
                    }
                    for wh in webhooks
                ],
            }
    return {}


def collect_kac_admission_latency() -> dict:
    """Parse apiserver /metrics for webhook latency specific to falcon-kac."""
    raw = kubectl_raw("/metrics")
    lat = {"p50": None, "p95": None, "p99": None, "sample_count": 0, "mean_ms": None}
    if not raw:
        return lat
    buckets, total_count, total_sum = {}, 0, 0.0
    for line in raw.splitlines():
        if "apiserver_admission_webhook_admission_duration_seconds" not in line:
            continue
        if WEBHOOK_PATTERN not in line.lower():
            continue
        if line.startswith("#"):
            continue
        if "_bucket{" in line:
            try:
                le_s = line.index('le="') + 4
                le_e = line.index('"', le_s)
                le_v = line[le_s:le_e]
                val  = float(line.split()[-1])
                if le_v != "+Inf":
                    buckets[float(le_v)] = val
            except (ValueError, IndexError):
                pass
        elif "_count{" in line:
            try:   total_count = int(float(line.split()[-1]))
            except (ValueError, IndexError): pass
        elif "_sum{" in line:
            try:   total_sum = float(line.split()[-1])
            except (ValueError, IndexError): pass
    if total_count > 0:
        lat["sample_count"] = total_count
        lat["mean_ms"] = round(total_sum / total_count * 1000, 2)
    if buckets and total_count > 0:
        for pct, key in [(0.5, "p50"), (0.95, "p95"), (0.99, "p99")]:
            target = pct * total_count
            for bound in sorted(buckets):
                if buckets[bound] >= target:
                    lat[key] = round(bound * 1000, 2)
                    break
    return lat


def collect_kac_hpa(namespace: str) -> dict:
    """Check if a HorizontalPodAutoscaler exists for falcon-kac."""
    hpa = kubectl_json(["get", "hpa", DEPLOY_NAME, "-n", namespace])
    if not hpa:
        return {}
    spec   = hpa.get("spec",   {})
    status = hpa.get("status", {})
    return {
        "min_replicas": spec.get("minReplicas", 1),
        "max_replicas": spec.get("maxReplicas", 1),
        "current":      status.get("currentReplicas", 0),
        "desired":      status.get("desiredReplicas", 0),
        "metrics":      [m.get("type") for m in spec.get("metrics", [])],
    }


def collect_resource_quota(namespace: str) -> dict:
    """Check ResourceQuota in the falcon-kac namespace."""
    rqs = kubectl_json(["get", "resourcequota", "-n", namespace])
    if not rqs or not rqs.get("items"):
        return {}
    out = {}
    for rq in rqs.get("items", []):
        name   = rq["metadata"]["name"]
        status = rq.get("status", {})
        out[name] = {
            "hard": status.get("hard", {}),
            "used": status.get("used", {}),
        }
    return out


def collect_events(namespace: str) -> list:
    evts = kubectl_json(["get", "events", "-n", namespace,
                         "--sort-by=.lastTimestamp"])
    if not evts:
        return []
    out = []
    for e in evts.get("items", [])[-20:]:
        out.append({
            "type":     e.get("type", "Normal"),
            "reason":   e.get("reason", ""),
            "message":  e.get("message", "")[:150],
            "count":    e.get("count", 1),
            "last_ts":  e.get("lastTimestamp", ""),
            "object":   e.get("involvedObject", {}).get("name", ""),
        })
    return [e for e in out if e["type"] == "Warning"] or out[-5:]


# ─────────────────────────────────────────────────────────────────────────────
# PARSING HELPERS
# ─────────────────────────────────────────────────────────────────────────────
def _parse_cpu(raw: str) -> float:
    raw = raw.strip()
    if raw.endswith("m"):
        return float(raw[:-1]) / 1000
    try:   return float(raw)
    except ValueError: return 0.0


def _parse_mem(raw: str) -> float:
    raw = raw.strip()
    if   raw.endswith("Ki"): return float(raw[:-2]) / 1024
    elif raw.endswith("Mi"): return float(raw[:-2])
    elif raw.endswith("Gi"): return float(raw[:-2]) * 1024
    elif raw.endswith("Ti"): return float(raw[:-2]) * 1024 * 1024
    try:   return float(raw) / (1024 * 1024)
    except ValueError: return 0.0


# ─────────────────────────────────────────────────────────────────────────────
# SNAPSHOT
# ─────────────────────────────────────────────────────────────────────────────
def collect_snapshot(label: str, namespace: str, release: str) -> dict:
    print(f"\n[{label.upper()}] {datetime.datetime.now().isoformat()}")
    snap = {
        "label":            label,
        "timestamp":        datetime.datetime.now().isoformat(),
        "namespace":        namespace,
        "release":          release,
        "cluster_info":     {},
        "helm_release":     {},
        "kac_deployment":   {},
        "kac_pods":         [],
        "webhook_config":   {},
        "admission_latency":{},
        "hpa":              {},
        "resource_quota":   {},
        "events":           [],
    }

    print("  → Cluster identity ...")
    snap["cluster_info"] = collect_cluster_info()
    ci = snap["cluster_info"]
    print(f"     Cluster : {ci['cluster_name']}  |  Platform: {ci['k8s_platform']}  |  K8s: {ci['k8s_version']}")

    print(f"  → Helm release [{release}] in namespace [{namespace}] ...")
    snap["helm_release"] = collect_helm_release(namespace, release)
    hr = snap["helm_release"]
    print(f"     Chart: {hr['chart']}-{hr['chart_version']}  |  Status: {hr['status']}  |  AppVersion: {hr['app_version']}")

    print("  → KAC Deployment status ...")
    snap["kac_deployment"] = collect_kac_deployment(namespace)

    print("  → KAC pod resources (top pods --containers) ...")
    snap["kac_pods"] = collect_kac_pods(namespace)

    print("  → ValidatingWebhookConfiguration ...")
    snap["webhook_config"] = collect_webhook_config()

    print("  → Admission webhook latency from apiserver /metrics ...")
    snap["admission_latency"] = collect_kac_admission_latency()

    print("  → HPA (if configured) ...")
    snap["hpa"] = collect_kac_hpa(namespace)

    print("  → ResourceQuota ...")
    snap["resource_quota"] = collect_resource_quota(namespace)

    print("  → Events ...")
    snap["events"] = collect_events(namespace)

    return snap


def collect_samples(n: int, interval_s: int, namespace: str, release: str) -> list:
    samples = []
    for i in range(n):
        print(f"\n── Sample {i+1}/{n} ──────────────────────────────────────────")
        samples.append(collect_snapshot(f"sample_{i+1}", namespace, release))
        if i < n - 1:
            print(f"  Sleeping {interval_s}s ...")
            time.sleep(interval_s)
    return samples


# ─────────────────────────────────────────────────────────────────────────────
# DELTA / ANALYSIS
# ─────────────────────────────────────────────────────────────────────────────
def compute_delta(baseline: dict, post: dict) -> dict:
    """Compute KAC-specific deltas."""
    delta = {"pods": [], "summary": {}}

    base_pods = {p["pod"]: p for p in baseline.get("kac_pods", [])}
    post_pods = {p["pod"]: p for p in post.get("kac_pods", [])}

    for pname, pp in post_pods.items():
        bp = base_pods.get(pname, {"cpu_cores": 0, "mem_mib": 0})
        delta["pods"].append({
            "pod":          pname,
            "node":         pp.get("node", ""),
            "cpu_base":     bp["cpu_cores"],
            "cpu_post":     pp["cpu_cores"],
            "cpu_delta":    round(pp["cpu_cores"] - bp["cpu_cores"], 4),
            "mem_base":     bp["mem_mib"],
            "mem_post":     pp["mem_mib"],
            "mem_delta":    round(pp["mem_mib"] - bp["mem_mib"], 2),
            "containers":   pp.get("containers", []),
        })

    base_lat = baseline.get("admission_latency", {})
    post_lat  = post.get("admission_latency", {})
    delta["latency_delta"] = {
        "p50": _safe_diff(post_lat.get("p50"), base_lat.get("p50")),
        "p95": _safe_diff(post_lat.get("p95"), base_lat.get("p95")),
        "p99": _safe_diff(post_lat.get("p99"), base_lat.get("p99")),
    }

    total_base_cpu = sum(p["cpu_cores"] for p in baseline.get("kac_pods", []))
    total_post_cpu = sum(p["cpu_cores"] for p in post.get("kac_pods",     []))
    total_base_mem = sum(p["mem_mib"]   for p in baseline.get("kac_pods", []))
    total_post_mem = sum(p["mem_mib"]   for p in post.get("kac_pods",     []))

    delta["summary"] = {
        "cpu_base":      round(total_base_cpu, 4),
        "cpu_post":      round(total_post_cpu, 4),
        "cpu_delta":     round(total_post_cpu - total_base_cpu, 4),
        "cpu_pct":       round((total_post_cpu - total_base_cpu) / max(total_base_cpu, 0.001) * 100, 1),
        "mem_base":      round(total_base_mem, 2),
        "mem_post":      round(total_post_mem, 2),
        "mem_delta":     round(total_post_mem - total_base_mem, 2),
        "mem_pct":       round((total_post_mem - total_base_mem) / max(total_base_mem, 0.001) * 100, 1),
        "pod_count":     len(post_pods),
    }
    return delta


def _safe_diff(a, b):
    if a is None or b is None: return None
    return round(a - b, 2)


def assess_risk(delta: dict, post: dict) -> dict:
    findings   = []
    risk_level = "LOW"
    s          = delta.get("summary", {})
    dep        = post.get("kac_deployment", {})
    lat        = post.get("admission_latency", {})

    # Deployment health
    desired, ready = dep.get("desired", 0), dep.get("ready", 0)
    unavail         = dep.get("unavailable", 0)
    if desired > 0 and ready < desired:
        risk_level = "HIGH"
        findings.append(f"CRIT  KAC Deployment: only {ready}/{desired} replicas ready ({unavail} unavailable).")
    elif desired > 0:
        findings.append(f"OK    KAC Deployment: {ready}/{desired} replicas ready.")
    if ready < 2:
        findings.append("WARN  Running fewer than 2 KAC replicas — no HA, a pod restart will cause downtime.")
        risk_level = "MEDIUM" if risk_level == "LOW" else risk_level

    # CPU
    cpu_pct = s.get("cpu_pct", 0)
    if cpu_pct > 80:
        risk_level = "HIGH"
        findings.append(f"CRIT  KAC CPU increased {cpu_pct}% — pods may be near CPU limits.")
    elif cpu_pct > 40:
        risk_level = "MEDIUM" if risk_level == "LOW" else risk_level
        findings.append(f"WARN  KAC CPU increased {cpu_pct}% — monitor throttling.")
    else:
        findings.append(f"OK    KAC CPU delta {cpu_pct}% — within normal range.")

    # Memory
    mem_pct = s.get("mem_pct", 0)
    if mem_pct > 50:
        risk_level = "HIGH"
        findings.append(f"CRIT  KAC memory increased {mem_pct}% — approaching limits.")
    elif mem_pct > 20:
        findings.append(f"WARN  KAC memory increased {mem_pct}% — review limits.")
        risk_level = "MEDIUM" if risk_level == "LOW" else risk_level
    else:
        findings.append(f"OK    KAC memory delta {mem_pct}% — within normal range.")

    # Webhook config
    wh_cfg = post.get("webhook_config", {})
    for wh in wh_cfg.get("webhooks", []):
        fp      = wh.get("failurePolicy", "Fail")
        timeout = wh.get("timeoutSeconds", 10)
        if fp == "Fail":
            findings.append(
                f"WARN  Webhook '{wh['name']}' failurePolicy=Fail / timeoutSeconds={timeout} — "
                "KAC outage blocks pod scheduling cluster-wide."
            )
            risk_level = "MEDIUM" if risk_level == "LOW" else risk_level
        else:
            findings.append(f"OK    Webhook '{wh['name']}' failurePolicy=Ignore — safe failure mode.")

    # Latency
    p99 = lat.get("p99")
    if p99 and p99 > 1000:
        risk_level = "HIGH"
        findings.append(f"CRIT  KAC p99 latency {p99}ms — exceeds 1000ms critical threshold.")
    elif p99 and p99 > 500:
        findings.append(f"WARN  KAC p99 latency {p99}ms — approaching 1000ms threshold.")
        risk_level = "MEDIUM" if risk_level == "LOW" else risk_level
    elif p99:
        findings.append(f"OK    KAC p99 latency {p99}ms — healthy.")

    # Restart count
    for pod in post.get("kac_pods", []):
        restarts = pod.get("restarts", 0)
        if restarts > 5:
            risk_level = "MEDIUM" if risk_level == "LOW" else risk_level
            findings.append(f"WARN  Pod {pod['pod']} has {restarts} restarts — check OOMKill or probe failures.")

    return {"risk_level": risk_level, "findings": findings}


# ─────────────────────────────────────────────────────────────────────────────
# CHARTS
# ─────────────────────────────────────────────────────────────────────────────
def generate_charts(post: dict, delta: dict, output_dir: str) -> list:
    if not HAS_PLOTLY:
        print("  [WARN] plotly not installed — skipping charts.")
        return []
    os.makedirs(output_dir, exist_ok=True)
    charts = []

    # 1. Per-pod CPU + memory (post-deploy)
    pods = post.get("kac_pods", [])
    if pods:
        pod_names = [p["pod"][:30] for p in pods]
        from plotly.subplots import make_subplots
        fig = make_subplots(rows=1, cols=2, subplot_titles=("CPU (cores)", "Memory (MiB)"))
        fig.add_trace(go.Bar(x=pod_names, y=[p["cpu_cores"] for p in pods], name="CPU"),    row=1, col=1)
        fig.add_trace(go.Bar(x=pod_names, y=[p["mem_mib"]   for p in pods], name="Memory"), row=1, col=2)
        fig.update_layout(
            title={"text": "KAC Pod Resource Usage (Post-Deploy)<br>"
                           "<span style='font-size:14px;font-weight:normal;'>"
                           "CPU cores and memory per falcon-kac pod</span>"},
            showlegend=False,
        )
        fig.update_xaxes(tickangle=-20)
        path = os.path.join(output_dir, "chart_kac_pods.png")
        fig.write_image(path)
        charts.append(path)

    # 2. Per-container breakdown (falcon-ac vs falcon-client)
    container_rows = []
    for pod in pods:
        for c in pod.get("containers", []):
            container_rows.append({
                "label": f"{pod['pod'][:20]}/{c['container']}",
                "cpu":   c["cpu_cores"],
                "mem":   c["mem_mib"],
            })
    if container_rows:
        labels  = [r["label"] for r in container_rows]
        cpu_v   = [r["cpu"]   for r in container_rows]
        mem_v   = [r["mem"]   for r in container_rows]
        fig = go.Figure()
        fig.add_trace(go.Bar(name="CPU (cores)", x=labels, y=cpu_v))
        fig.add_trace(go.Bar(name="Mem (MiB)",   x=labels, y=mem_v))
        fig.update_layout(
            title={"text": "KAC Container-Level Resource Breakdown<br>"
                           "<span style='font-size:14px;font-weight:normal;'>"
                           "falcon-ac vs falcon-client per pod</span>"},
            barmode="group",
            legend=dict(orientation="h", yanchor="bottom", y=1.05, xanchor="center", x=0.5),
        )
        fig.update_xaxes(title_text="Pod / Container", tickangle=-20)
        fig.update_yaxes(title_text="Usage")
        path = os.path.join(output_dir, "chart_kac_containers.png")
        fig.write_image(path)
        charts.append(path)

    # 3. Webhook latency percentiles
    lat = post.get("admission_latency", {})
    lat_vals = {k: v for k, v in lat.items() if k in ("p50","p95","p99") and v is not None}
    if lat_vals:
        fig = go.Figure(go.Bar(
            x=list(lat_vals.keys()),
            y=list(lat_vals.values()),
            text=[f"{v} ms" for v in lat_vals.values()],
            textposition="outside",
        ))
        fig.add_hline(y=500,  line_dash="dash", line_color="#ff9900",
                      annotation_text="500ms warn",  annotation_position="top right")
        fig.add_hline(y=1000, line_dash="dot",  line_color="#cc0000",
                      annotation_text="1000ms crit", annotation_position="top right")
        fig.update_layout(
            title={"text": "KAC Admission Webhook Latency<br>"
                           "<span style='font-size:14px;font-weight:normal;'>"
                           "p50 / p95 / p99 milliseconds from apiserver /metrics</span>"},
        )
        fig.update_xaxes(title_text="Percentile")
        fig.update_yaxes(title_text="Latency (ms)")
        path = os.path.join(output_dir, "chart_kac_latency.png")
        fig.write_image(path)
        charts.append(path)

    # 4. CPU baseline vs post (pod delta)
    dpods = delta.get("pods", [])
    if dpods:
        dnames = [p["pod"][:28] for p in dpods]
        fig = go.Figure()
        fig.add_trace(go.Bar(name="Baseline",    x=dnames, y=[p["cpu_base"] for p in dpods]))
        fig.add_trace(go.Bar(name="Post-Deploy", x=dnames, y=[p["cpu_post"] for p in dpods]))
        fig.update_layout(
            title={"text": "KAC Pod CPU: Baseline vs Post-Deploy<br>"
                           "<span style='font-size:14px;font-weight:normal;'>"
                           "CPU cores per pod</span>"},
            barmode="group",
            legend=dict(orientation="h", yanchor="bottom", y=1.05, xanchor="center", x=0.5),
        )
        fig.update_xaxes(title_text="Pod")
        fig.update_yaxes(title_text="CPU (cores)")
        path = os.path.join(output_dir, "chart_kac_cpu_delta.png")
        fig.write_image(path)
        charts.append(path)

    return charts


# ─────────────────────────────────────────────────────────────────────────────
# DOCX REPORT
# ─────────────────────────────────────────────────────────────────────────────
def _cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _add_table(doc, headers, rows, hdr_fill="1F3864", hdr_font="FFFFFF"):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(
            int(hdr_font[0:2],16), int(hdr_font[2:4],16), int(hdr_font[4:6],16))
        _cell_bg(hdr_cells[i], hdr_fill)
    for ridx, row in enumerate(rows):
        rc = tbl.add_row().cells
        fill = "EEF0F5" if ridx % 2 == 0 else "FFFFFF"
        for cidx, val in enumerate(row):
            rc[cidx].text = str(val)
            _cell_bg(rc[cidx], fill)
    doc.add_paragraph()


def generate_docx_report(baseline, post, delta, assessment, charts, output_path):
    if not HAS_DOCX:
        print("[ERROR] python-docx not installed — skipping DOCX."); return

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    ci  = post.get("cluster_info", {})
    hr  = post.get("helm_release", {})
    dep = post.get("kac_deployment", {})
    lat = post.get("admission_latency", {})
    wh  = post.get("webhook_config", {})
    s   = delta.get("summary", {})

    # Title
    t = doc.add_heading("CrowdStrike Falcon KAC", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("Performance Impact Assessment Report")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(14); p.runs[0].bold = True
    doc.add_paragraph()

    # 1. Cluster + Helm Identity
    doc.add_heading("1. Cluster & Helm Release Identity", level=1)
    _add_table(doc,
        headers=["Field", "Value"],
        rows=[
            ["Cluster Name",          ci.get("cluster_name",  "unknown")],
            ["Kubeconfig Context",    ci.get("context",        "unknown")],
            ["API Server URL",        ci.get("server_url",     "unknown")],
            ["Platform",              ci.get("k8s_platform",   "unknown")],
            ["Kubernetes Version",    ci.get("k8s_version",    "unknown")],
            ["Total / Worker Nodes",  f"{ci.get('node_count','N/A')} / {ci.get('worker_count','N/A')}"],
            ["Helm Release Name",     hr.get("release",        "unknown")],
            ["Helm Namespace",        hr.get("namespace",      "unknown")],
            ["Chart",                 f"{hr.get('chart','?')}-{hr.get('chart_version','?')}"],
            ["App Version",           hr.get("app_version",    "unknown")],
            ["Helm Status",           hr.get("status",         "unknown")],
            ["Last Deployed",         hr.get("deployed_at",    "unknown")],
            ["Baseline Captured",     baseline.get("timestamp","N/A")],
            ["Post-Deploy Captured",  post.get("timestamp",    "N/A")],
            ["Report Generated",      datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")],
        ],
    )

    # 2. Risk Assessment
    doc.add_heading("2. Risk Assessment", level=1)
    risk = assessment["risk_level"]
    risk_rgb = {"LOW": (26,127,55), "MEDIUM": (154,103,0), "HIGH": (185,28,28)}
    p = doc.add_paragraph()
    run = p.add_run(f"Overall Risk Level:  {risk}")
    run.bold = True; run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(*risk_rgb.get(risk, (0,0,0)))
    for f in assessment["findings"]:
        doc.add_paragraph(f, style="List Bullet")
    doc.add_paragraph()

    # 3. Resource Delta
    doc.add_heading("3. KAC Resource Delta Summary", level=1)
    _add_table(doc,
        headers=["Metric", "Baseline", "Post-Deploy", "Delta", "% Change"],
        rows=[
            ["KAC Total CPU (cores)",
             f"{s.get('cpu_base',0):.3f}", f"{s.get('cpu_post',0):.3f}",
             f"{s.get('cpu_delta',0):+.3f}", f"{s.get('cpu_pct',0):+.1f}%"],
            ["KAC Total Memory (MiB)",
             f"{s.get('mem_base',0):.0f}", f"{s.get('mem_post',0):.0f}",
             f"{s.get('mem_delta',0):+.0f}", f"{s.get('mem_pct',0):+.1f}%"],
            ["Pod Count", "N/A", str(s.get("pod_count",0)), "–", "–"],
        ],
    )

    # 4. Per-pod delta
    doc.add_heading("4. Per-Pod Resource Detail", level=1)
    _add_table(doc,
        headers=["Pod", "Node", "CPU Base", "CPU Post", "CPU Δ", "Mem Base", "Mem Post", "Mem Δ"],
        rows=[
            [p["pod"], p["node"],
             f"{p['cpu_base']:.3f}", f"{p['cpu_post']:.3f}", f"{p['cpu_delta']:+.3f}",
             f"{p['mem_base']:.0f}", f"{p['mem_post']:.0f}", f"{p['mem_delta']:+.0f}"]
            for p in delta.get("pods", [])
        ],
    )

    # 5. Per-container breakdown
    doc.add_heading("5. Per-Container Breakdown (Post-Deploy)", level=1)
    container_rows = []
    for pod in post.get("kac_pods", []):
        for c in pod.get("containers", []):
            container_rows.append([
                pod["pod"], c["container"],
                f"{c['cpu_cores']:.3f}", f"{c['mem_mib']:.0f}",
                str(pod.get("restarts", 0)), pod.get("node",""),
            ])
    if container_rows:
        _add_table(doc,
            headers=["Pod", "Container", "CPU (cores)", "Memory (MiB)", "Restarts", "Node"],
            rows=container_rows,
        )
    else:
        doc.add_paragraph("Container-level metrics not available (kubectl top --containers may need metrics-server).")

    # 6. Deployment Status
    doc.add_heading("6. KAC Deployment Status", level=1)
    cond_str = ", ".join(f"{c['type']}={c['status']}" for c in dep.get("conditions", []))
    _add_table(doc,
        headers=["Field", "Value"],
        rows=[
            ["Namespace",    dep.get("namespace",  "")],
            ["Desired",      str(dep.get("desired",  ""))],
            ["Ready",        str(dep.get("ready",    ""))],
            ["Available",    str(dep.get("available",""))],
            ["Unavailable",  str(dep.get("unavailable",""))],
            ["Conditions",   cond_str],
        ],
    )

    # Per-container resource requests/limits
    doc.add_heading("6.1  Container Resource Requests / Limits", level=2)
    cr_rows = []
    for cname, res in dep.get("container_resources", {}).items():
        reqs = res.get("requests", {})
        lims = res.get("limits",   {})
        cr_rows.append([
            cname,
            reqs.get("cpu",    "–"), reqs.get("memory", "–"),
            lims.get("cpu",    "–"), lims.get("memory",  "–"),
        ])
    if cr_rows:
        _add_table(doc,
            headers=["Container", "Req CPU", "Req Memory", "Limit CPU", "Limit Memory"],
            rows=cr_rows,
        )

    # 7. Helm values summary
    vs = hr.get("values_summary", {})
    if vs:
        doc.add_heading("7. Helm Values (Performance-Relevant)", level=1)
        _add_table(doc,
            headers=["Parameter", "Value"],
            rows=[
                ["replicas",          str(vs.get("replicas",          ""))],
                ["webhook.failurePolicy",  str(vs.get("failurePolicy", ""))],
                ["webhook.timeoutSeconds", str(vs.get("timeoutSeconds",""))],
                ["priorityClassName", str(vs.get("priorityClassName", ""))],
                ["resources.falcon-ac",    str(vs.get("resources_ac", ""))],
                ["resources.falcon-client",str(vs.get("resources_client",""))],
            ],
        )

    # 8. Webhook Config
    doc.add_heading("8. ValidatingWebhookConfiguration", level=1)
    if wh:
        doc.add_paragraph(f"Name: {wh.get('name','')}")
        _add_table(doc,
            headers=["Webhook", "Failure Policy", "Timeout (s)", "Match Policy",
                     "SideEffects", "NS Selector", "Rules"],
            rows=[
                [h["name"], h["failurePolicy"], str(h["timeoutSeconds"]),
                 h["matchPolicy"], h["sideEffects"],
                 "Yes" if h["namespaceSelector"] else "No", str(h["rules"])]
                for h in wh.get("webhooks", [])
            ],
        )
    else:
        doc.add_paragraph("No Falcon ValidatingWebhookConfiguration found.")

    # 9. Admission latency
    doc.add_heading("9. Admission Webhook Latency", level=1)
    lat_rows = []
    for k, lbl in [("p50","p50 (ms)"),("p95","p95 (ms)"),("p99","p99 (ms)"),("mean_ms","Mean (ms)")]:
        v = lat.get(k)
        if v is not None:
            lat_rows.append([lbl, str(v)])
    lat_rows.append(["Sample Count", str(lat.get("sample_count",0))])
    _add_table(doc, headers=["Metric","Value"], rows=lat_rows)

    # 10. HPA
    hpa = post.get("hpa", {})
    if hpa:
        doc.add_heading("10. Horizontal Pod Autoscaler", level=1)
        _add_table(doc,
            headers=["Field","Value"],
            rows=[
                ["Min Replicas",  str(hpa.get("min_replicas",""))],
                ["Max Replicas",  str(hpa.get("max_replicas",""))],
                ["Current",       str(hpa.get("current",""))],
                ["Desired",       str(hpa.get("desired",""))],
                ["Metrics",       ", ".join(hpa.get("metrics",[]))],
            ],
        )

    # 11. ResourceQuota
    rq = post.get("resource_quota", {})
    if rq:
        doc.add_heading("11. ResourceQuota", level=1)
        rq_rows = []
        for qname, qdata in rq.items():
            for resource, hard_val in qdata.get("hard", {}).items():
                used_val = qdata.get("used", {}).get(resource, "–")
                rq_rows.append([qname, resource, hard_val, used_val])
        if rq_rows:
            _add_table(doc, headers=["Quota Name","Resource","Hard Limit","Used"], rows=rq_rows)

    # 12. Events
    events = post.get("events", [])
    if events:
        doc.add_heading("12. Events (falcon-kac namespace)", level=1)
        _add_table(doc,
            headers=["Type","Reason","Count","Last Seen","Object","Message"],
            rows=[
                [e["type"], e["reason"], str(e["count"]),
                 e["last_ts"], e["object"], e["message"]]
                for e in events
            ],
        )

    # 13. Charts
    if charts:
        doc.add_heading("13. Performance Charts", level=1)
        titles = {
            "chart_kac_pods.png":        "Figure 1 – KAC Pod Resource Usage (Post-Deploy)",
            "chart_kac_containers.png":  "Figure 2 – KAC Container-Level Breakdown (falcon-ac vs falcon-client)",
            "chart_kac_latency.png":     "Figure 3 – Admission Webhook Latency p50/p95/p99",
            "chart_kac_cpu_delta.png":   "Figure 4 – KAC Pod CPU: Baseline vs Post-Deploy",
        }
        for cp in charts:
            if os.path.exists(cp):
                doc.add_paragraph(titles.get(os.path.basename(cp), os.path.basename(cp))).runs[0].bold = True
                doc.add_picture(cp, width=Inches(6.0))
                doc.add_paragraph()

    # 14. Recommendations
    doc.add_heading("14. Recommended Next Steps", level=1)
    for i, rec in enumerate([
        "Scale to ≥2 replicas for HA: helm upgrade falcon-kac ... --set replicas=2",
        "If p99 > 500ms, increase falcon-ac CPU limits in values.yaml under resources.",
        "Set webhook.failurePolicy=Ignore with webhook.timeoutSeconds=10 for availability.",
        "Add priorityClassName: system-cluster-critical in values.yaml to prevent eviction.",
        "Enable HPA targeting CPU utilisation at 60% to auto-scale under burst load.",
        "Monitor container_cpu_cfs_throttled_periods_total{container='falcon-ac'} for throttling.",
        "Re-run during peak load for representative admission latency measurements.",
    ], 1):
        doc.add_paragraph(rec, style="List Number")

    # Footer
    doc.add_paragraph()
    fp = doc.add_paragraph(
        f"falcon_kac_perf_assessment.py  |  "
        f"Cluster: {ci.get('cluster_name','unknown')}  |  "
        f"Release: {hr.get('release','unknown')} ({hr.get('chart_version','?')})  |  "
        f"{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )
    fp.runs[0].font.size = Pt(9)
    fp.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(output_path)
    print(f"✅  DOCX report written to: {output_path}")


# ─────────────────────────────────────────────────────────────────────────────
# HTML REPORT
# ─────────────────────────────────────────────────────────────────────────────
HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>Falcon KAC – Performance Impact Report</title>
<style>
  body { font-family:'Segoe UI',Arial,sans-serif; background:#0f1117; color:#e0e0e0; margin:40px; }
  h1   { color:#e85c27; font-size:2em; }
  h2   { color:#ff7b4f; border-bottom:1px solid #333; padding-bottom:6px; margin-top:28px; }
  h3   { color:#c0c0c0; margin-top:16px; }
  table{ width:100%; border-collapse:collapse; margin-bottom:20px; }
  th   { background:#1e2230; color:#ff7b4f; padding:8px 12px; text-align:left; }
  td   { padding:7px 12px; border-bottom:1px solid #2a2d3a; }
  tr:nth-child(even) td { background:#181c29; }
  .badge-low    { background:#1a7f37; color:#fff; padding:2px 12px; border-radius:12px; font-weight:bold; }
  .badge-medium { background:#9a6700; color:#fff; padding:2px 12px; border-radius:12px; font-weight:bold; }
  .badge-high   { background:#b91c1c; color:#fff; padding:2px 12px; border-radius:12px; font-weight:bold; }
  .finding { padding:6px 0; border-bottom:1px dashed #2a2d3a; font-family:monospace; }
  .chart   { max-width:100%; margin:16px 0; border:1px solid #333; border-radius:8px; }
  .section { background:#13161f; border-radius:8px; padding:16px 24px; margin-bottom:24px; }
  .meta    { color:#888; font-size:0.85em; }
  code     { background:#1a1d2b; padding:2px 6px; border-radius:4px; font-family:monospace; }
  .ok   { color:#4caf50; } .warn { color:#ff9800; } .crit { color:#f44336; }
</style>
</head>
<body>
<h1>🛡️ CrowdStrike Falcon KAC – Performance Impact Report</h1>
<p class="meta">Generated: {{ timestamp }}</p>

<div class="section">
  <h2>Cluster &amp; Helm Release Identity</h2>
  <table>
    <tr><th>Field</th><th>Value</th></tr>
    <tr><td>Cluster Name</td><td><strong>{{ ci.cluster_name }}</strong></td></tr>
    <tr><td>Kubeconfig Context</td><td>{{ ci.context }}</td></tr>
    <tr><td>API Server URL</td><td><code>{{ ci.server_url }}</code></td></tr>
    <tr><td>Platform</td><td>{{ ci.k8s_platform }}</td></tr>
    <tr><td>Kubernetes Version</td><td>{{ ci.k8s_version }}</td></tr>
    <tr><td>Nodes (Total / Workers)</td><td>{{ ci.node_count | default('N/A') }} / {{ ci.worker_count | default('N/A') }}</td></tr>
    <tr><td>Helm Release</td><td><strong>{{ hr.release }}</strong></td></tr>
    <tr><td>Helm Namespace</td><td>{{ hr.namespace }}</td></tr>
    <tr><td>Chart</td><td>{{ hr.chart }}-{{ hr.chart_version }}</td></tr>
    <tr><td>App Version</td><td>{{ hr.app_version }}</td></tr>
    <tr><td>Helm Status</td><td>{{ hr.status }}</td></tr>
    <tr><td>Last Deployed</td><td>{{ hr.deployed_at }}</td></tr>
    <tr><td>Baseline Captured</td><td>{{ baseline_ts }}</td></tr>
    <tr><td>Post-Deploy Captured</td><td>{{ post_ts }}</td></tr>
  </table>
</div>

<div class="section">
  <h2>Risk Assessment</h2>
  <p>Risk Level: <span class="badge-{{ risk_level | lower }}">{{ risk_level }}</span></p>
  {% for f in findings %}
  <div class="finding
    {%- if f.startswith('CRIT') %} crit
    {%- elif f.startswith('WARN') %} warn
    {%- else %} ok{% endif %}">{{ f }}</div>
  {% endfor %}
</div>

<div class="section">
  <h2>KAC Resource Delta</h2>
  <table>
    <tr><th>Metric</th><th>Baseline</th><th>Post-Deploy</th><th>Delta</th><th>% Change</th></tr>
    <tr>
      <td>KAC Total CPU (cores)</td>
      <td>{{ "%.3f"|format(s.cpu_base) }}</td><td>{{ "%.3f"|format(s.cpu_post) }}</td>
      <td>{{ "%+.3f"|format(s.cpu_delta) }}</td><td>{{ "%+.1f"|format(s.cpu_pct) }}%</td>
    </tr>
    <tr>
      <td>KAC Total Memory (MiB)</td>
      <td>{{ "%.0f"|format(s.mem_base) }}</td><td>{{ "%.0f"|format(s.mem_post) }}</td>
      <td>{{ "%+.0f"|format(s.mem_delta) }}</td><td>{{ "%+.1f"|format(s.mem_pct) }}%</td>
    </tr>
    <tr><td>Pod Count</td><td>N/A</td><td>{{ s.pod_count }}</td><td>–</td><td>–</td></tr>
  </table>
</div>

<div class="section">
  <h2>Per-Pod Detail</h2>
  <table>
    <tr><th>Pod</th><th>Node</th><th>CPU Base</th><th>CPU Post</th><th>CPU Δ</th>
        <th>Mem Base (MiB)</th><th>Mem Post (MiB)</th><th>Mem Δ (MiB)</th></tr>
    {% for p in pod_delta %}
    <tr>
      <td>{{ p.pod }}</td><td>{{ p.node }}</td>
      <td>{{ "%.3f"|format(p.cpu_base) }}</td><td>{{ "%.3f"|format(p.cpu_post) }}</td>
      <td>{{ "%+.3f"|format(p.cpu_delta) }}</td>
      <td>{{ "%.0f"|format(p.mem_base) }}</td><td>{{ "%.0f"|format(p.mem_post) }}</td>
      <td>{{ "%+.0f"|format(p.mem_delta) }}</td>
    </tr>
    {% endfor %}
  </table>
</div>

<div class="section">
  <h2>Per-Container Breakdown (Post-Deploy)</h2>
  <table>
    <tr><th>Pod</th><th>Container</th><th>CPU (cores)</th><th>Memory (MiB)</th><th>Restarts</th><th>Node</th></tr>
    {% for pod in kac_pods %}
      {% for c in pod.containers %}
      <tr>
        <td>{{ pod.pod }}</td><td>{{ c.container }}</td>
        <td>{{ "%.3f"|format(c.cpu_cores) }}</td><td>{{ "%.0f"|format(c.mem_mib) }}</td>
        <td>{{ pod.restarts }}</td><td>{{ pod.node }}</td>
      </tr>
      {% endfor %}
    {% endfor %}
  </table>
</div>

<div class="section">
  <h2>KAC Deployment Status</h2>
  {% if dep %}
  <table>
    <tr><th>Namespace</th><th>Desired</th><th>Ready</th><th>Available</th><th>Unavailable</th></tr>
    <tr><td>{{ dep.namespace }}</td><td>{{ dep.desired }}</td><td>{{ dep.ready }}</td>
        <td>{{ dep.available }}</td><td>{{ dep.unavailable }}</td></tr>
  </table>
  <h3>Container Resource Requests / Limits</h3>
  <table>
    <tr><th>Container</th><th>Req CPU</th><th>Req Memory</th><th>Limit CPU</th><th>Limit Memory</th></tr>
    {% for cname, res in dep.container_resources.items() %}
    <tr>
      <td>{{ cname }}</td>
      <td>{{ res.requests.get('cpu','–') }}</td><td>{{ res.requests.get('memory','–') }}</td>
      <td>{{ res.limits.get('cpu','–') }}</td><td>{{ res.limits.get('memory','–') }}</td>
    </tr>
    {% endfor %}
  </table>
  {% endif %}
</div>

{% if hr.values_summary %}
<div class="section">
  <h2>Helm Values (Performance-Relevant)</h2>
  <table>
    <tr><th>Parameter</th><th>Value</th></tr>
    <tr><td>replicas</td><td>{{ hr.values_summary.replicas }}</td></tr>
    <tr><td>webhook.failurePolicy</td><td>{{ hr.values_summary.failurePolicy }}</td></tr>
    <tr><td>webhook.timeoutSeconds</td><td>{{ hr.values_summary.timeoutSeconds }}</td></tr>
    <tr><td>priorityClassName</td><td>{{ hr.values_summary.priorityClassName }}</td></tr>
    <tr><td>resources.falcon-ac</td><td><code>{{ hr.values_summary.resources_ac }}</code></td></tr>
    <tr><td>resources.falcon-client</td><td><code>{{ hr.values_summary.resources_client }}</code></td></tr>
  </table>
</div>
{% endif %}

<div class="section">
  <h2>Webhook Configuration</h2>
  {% if webhook %}
  <p>Name: <code>{{ webhook.name }}</code></p>
  <table>
    <tr><th>Webhook</th><th>Failure Policy</th><th>Timeout (s)</th><th>Match Policy</th>
        <th>SideEffects</th><th>NS Selector</th><th>Rules</th></tr>
    {% for wh in webhook.webhooks %}
    <tr><td>{{ wh.name }}</td><td>{{ wh.failurePolicy }}</td><td>{{ wh.timeoutSeconds }}</td>
        <td>{{ wh.matchPolicy }}</td><td>{{ wh.sideEffects }}</td>
        <td>{{ "Yes" if wh.namespaceSelector else "No" }}</td><td>{{ wh.rules }}</td></tr>
    {% endfor %}
  </table>
  {% else %}<p class="meta">No Falcon webhook found.</p>{% endif %}
</div>

<div class="section">
  <h2>Admission Webhook Latency</h2>
  <table>
    <tr><th>Metric</th><th>Value</th></tr>
    {% if lat.p50   is not none %}<tr><td>p50 (ms)</td><td>{{ lat.p50 }}</td></tr>{% endif %}
    {% if lat.p95   is not none %}<tr><td>p95 (ms)</td><td>{{ lat.p95 }}</td></tr>{% endif %}
    {% if lat.p99   is not none %}<tr><td>p99 (ms)</td><td>{{ lat.p99 }}</td></tr>{% endif %}
    {% if lat.mean_ms is not none %}<tr><td>Mean (ms)</td><td>{{ lat.mean_ms }}</td></tr>{% endif %}
    <tr><td>Sample Count</td><td>{{ lat.sample_count }}</td></tr>
  </table>
</div>

{% if hpa %}
<div class="section">
  <h2>Horizontal Pod Autoscaler</h2>
  <table>
    <tr><th>Min</th><th>Max</th><th>Current</th><th>Desired</th><th>Metrics</th></tr>
    <tr><td>{{ hpa.min_replicas }}</td><td>{{ hpa.max_replicas }}</td>
        <td>{{ hpa.current }}</td><td>{{ hpa.desired }}</td>
        <td>{{ hpa.metrics | join(', ') }}</td></tr>
  </table>
</div>
{% endif %}

{% if events %}
<div class="section">
  <h2>Events (falcon-kac namespace)</h2>
  <table>
    <tr><th>Type</th><th>Reason</th><th>Count</th><th>Last Seen</th><th>Object</th><th>Message</th></tr>
    {% for e in events %}
    <tr><td>{{ e.type }}</td><td>{{ e.reason }}</td><td>{{ e.count }}</td>
        <td>{{ e.last_ts }}</td><td>{{ e.object }}</td><td>{{ e.message }}</td></tr>
    {% endfor %}
  </table>
</div>
{% endif %}

{% if charts %}
<div class="section">
  <h2>Performance Charts</h2>
  {% for c in charts %}<img src="{{ c }}" class="chart"><br>{% endfor %}
</div>
{% endif %}

<div class="section">
  <h2>Recommended Next Steps</h2>
  <ol>
    <li>Scale to ≥2 replicas for HA: <code>helm upgrade falcon-kac ... --set replicas=2</code></li>
    <li>If p99 &gt; 500ms, increase <code>falcon-ac</code> CPU limits in values.yaml.</li>
    <li>Set <code>webhook.failurePolicy=Ignore</code> and <code>timeoutSeconds=10</code> for availability.</li>
    <li>Add <code>priorityClassName: system-cluster-critical</code> to prevent eviction.</li>
    <li>Enable HPA targeting 60% CPU to auto-scale under admission burst.</li>
    <li>Monitor <code>container_cpu_cfs_throttled_periods_total{container="falcon-ac"}</code> for throttling.</li>
    <li>Re-run during peak load for representative admission latency measurements.</li>
  </ol>
</div>
<p class="meta" style="text-align:center;">
  falcon_kac_perf_assessment.py &nbsp;|&nbsp;
  Cluster: {{ ci.cluster_name }} &nbsp;|&nbsp;
  Release: {{ hr.release }} ({{ hr.chart_version }}) &nbsp;|&nbsp; {{ timestamp }}
</p>
</body></html>"""


def generate_html_report(baseline, post, delta, assessment, charts, output_path):
    if not HAS_JINJA:
        print("[ERROR] jinja2 not installed — skipping HTML."); return
    report_dir = os.path.dirname(os.path.abspath(output_path))
    rel_charts  = [os.path.relpath(c, report_dir) for c in charts]
    html = Template(HTML_TEMPLATE).render(
        timestamp   = datetime.datetime.now().isoformat(),
        baseline_ts = baseline.get("timestamp", ""),
        post_ts     = post.get("timestamp",     ""),
        ci          = post.get("cluster_info",  {}),
        hr          = post.get("helm_release",  {}),
        risk_level  = assessment["risk_level"],
        findings    = assessment["findings"],
        s           = delta.get("summary",      {}),
        pod_delta   = delta.get("pods",         []),
        dep         = post.get("kac_deployment",{}),
        kac_pods    = post.get("kac_pods",      []),
        lat         = post.get("admission_latency", {}),
        webhook     = post.get("webhook_config",{}),
        hpa         = post.get("hpa",           {}),
        events      = post.get("events",        []),
        charts      = rel_charts,
    )
    Path(output_path).write_text(html, encoding="utf-8")
    print(f"✅  HTML report written to: {output_path}")


# ─────────────────────────────────────────────────────────────────────────────
# CLI
# ─────────────────────────────────────────────────────────────────────────────
def parse_args():
    p = argparse.ArgumentParser(
        description="Falcon KAC (Helm) Performance Impact Assessment",
        formatter_class=argparse.RawTextHelpFormatter,
        epilog="""
Examples:
  python3 falcon_kac_perf_assessment.py --mode baseline --output /tmp/kac_baseline.json

  python3 falcon_kac_perf_assessment.py --mode post \\
      --baseline /tmp/kac_baseline.json \\
      --output /tmp/kac_report \\
      --namespace falcon-kac \\
      --release falcon-kac

  python3 falcon_kac_perf_assessment.py --mode sample \\
      --samples 12 --interval 60 --output /tmp/kac_samples.json
        """,
    )
    p.add_argument("--mode",      required=True, choices=["baseline","post","sample"])
    p.add_argument("--baseline",  default=None)
    p.add_argument("--output",    required=True,
                   help="For 'post' mode: base path without extension (writes .html + .docx)")
    p.add_argument("--namespace", default=DEFAULT_NAMESPACE,
                   help=f"Helm / KAC namespace (default: {DEFAULT_NAMESPACE})")
    p.add_argument("--release",   default=DEFAULT_RELEASE,
                   help=f"Helm release name (default: {DEFAULT_RELEASE})")
    p.add_argument("--samples",   type=int, default=6)
    p.add_argument("--interval",  type=int, default=60)
    p.add_argument("--chart-dir", default=None)
    return p.parse_args()


def main():
    args = parse_args()
    ns   = args.namespace
    rel  = args.release

    if args.mode == "baseline":
        snap = collect_snapshot("baseline", ns, rel)
        Path(args.output).write_text(json.dumps(snap, indent=2))
        print(f"\n✅  Baseline saved to: {args.output}")

    elif args.mode == "sample":
        samples = collect_samples(args.samples, args.interval, ns, rel)
        Path(args.output).write_text(json.dumps(samples, indent=2))
        print(f"\n✅  {len(samples)} samples saved to: {args.output}")

    elif args.mode == "post":
        if not args.baseline:
            print("[ERROR] --baseline required for --mode post"); sys.exit(1)
        bpath = Path(args.baseline)
        if not bpath.exists():
            print(f"[ERROR] Baseline not found: {args.baseline}"); sys.exit(1)

        baseline = json.loads(bpath.read_text())
        post     = collect_snapshot("post", ns, rel)
        delta    = compute_delta(baseline, post)
        assess   = assess_risk(delta, post)

        base_out  = str(args.output).removesuffix(".html").removesuffix(".docx")
        chart_dir = args.chart_dir or str(Path(base_out).parent)
        charts    = generate_charts(post, delta, chart_dir)

        generate_html_report(baseline, post, delta, assess, charts, base_out + ".html")
        generate_docx_report(baseline, post, delta, assess, charts, base_out + ".docx")

        print("\n── Summary ────────────────────────────────────────────────")
        ci = post.get("cluster_info",{})
        hr = post.get("helm_release",{})
        print(f"  Cluster   : {ci.get('cluster_name')} ({ci.get('k8s_platform')})")
        print(f"  Release   : {hr.get('release')} chart={hr.get('chart_version')} status={hr.get('status')}")
        print(f"  Risk      : {assess['risk_level']}")
        for f in assess["findings"]:
            print(f"  {f}")
        print(f"  HTML → {base_out}.html")
        print(f"  DOCX → {base_out}.docx")
        print("───────────────────────────────────────────────────────────")


if __name__ == "__main__":
    main()
